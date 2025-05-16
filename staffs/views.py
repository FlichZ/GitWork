import csv
import sqlite3
from datetime import datetime
import json
import mimetypes
from django.conf import settings
from django.db import connections
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required, user_passes_test
from django.utils import timezone
from django.utils.translation import gettext_lazy as _
from django.db.models import Q
from .models import Document, Organization, User, UserActionLog, Notification, ChatMessage, Chat, DocumentTemplate, DocumentVersion, DigitalSignature
from .forms import SendDocumentForm, CustomUserCreationForm, OrganizationCreationForm, OrganizationEditForm
from django.contrib import messages as django_messages
from django.http import FileResponse, HttpResponse, JsonResponse
import os
import io
import base64
from pdf2image import convert_from_path
from docx import Document as DocxDocument
from PIL import Image, ImageDraw, ImageFont
import pandas as pd
import tempfile
import shutil
import logging
import subprocess
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side
from io import BytesIO
from django.views.decorators.http import require_POST, require_GET
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from PyPDF2 import PdfReader
from django.core.files import File
from django.urls import reverse

try:
    from docx2pdf import convert as docx2pdf_convert
    DOCX2PDF_AVAILABLE = True
except ImportError:
    DOCX2PDF_AVAILABLE = False

try:
    import unoconv
    UNOCONV_AVAILABLE = True
except ImportError:
    UNOCONV_AVAILABLE = False

try:
    from docx2python import docx2python
    DOCX2PYTHON_AVAILABLE = True
except ImportError:
    DOCX2PYTHON_AVAILABLE = False

# Проверяем доступность LibreOffice/OpenOffice
def check_libreoffice_available():
    try:
        result = subprocess.run(['which', 'soffice'], stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        return result.returncode == 0
    except Exception:
        return False

LIBREOFFICE_AVAILABLE = check_libreoffice_available()

# Функция для конвертации DOCX в PDF с сохранением форматирования
def convert_docx_to_pdf(input_path, output_path):
    """
    Конвертирует DOCX в PDF, пытаясь использовать различные методы, сохраняя форматирование.
    Возвращает True, если конвертация успешна, иначе False.
    """
    conversion_success = False
    error_message = ""
    
    # Способ 1: Использование docx2pdf (MS Word)
    if DOCX2PDF_AVAILABLE and not conversion_success:
        try:
            logger.info(f"Попытка конвертации с помощью docx2pdf: {input_path}")
            docx2pdf_convert(input_path, output_path)
            conversion_success = os.path.exists(output_path) and os.path.getsize(output_path) > 0
            if conversion_success:
                logger.info("Конвертация с помощью docx2pdf успешна")
        except Exception as e:
            error_message += f"docx2pdf: {str(e)}; "
            logger.warning(f"docx2pdf не смог конвертировать файл: {str(e)}")
    
    # Способ 2: Использование LibreOffice напрямую
    if LIBREOFFICE_AVAILABLE and not conversion_success:
        try:
            logger.info(f"Попытка конвертации с помощью LibreOffice: {input_path}")
            # Получаем абсолютные пути для корректной работы soffice
            abs_input = os.path.abspath(input_path)
            abs_output_dir = os.path.dirname(os.path.abspath(output_path))
            
            # Используем headless режим LibreOffice для конвертации
            cmd = [
                'soffice', '--headless', '--convert-to', 'pdf', 
                '--outdir', abs_output_dir, abs_input
            ]
            
            process = subprocess.run(
                cmd, 
                stdout=subprocess.PIPE, 
                stderr=subprocess.PIPE,
                timeout=30  # Таймаут 30 секунд
            )
            
            # Проверяем результат
            if process.returncode == 0:
                # LibreOffice сохраняет с тем же именем, но расширением .pdf
                base_name = os.path.basename(abs_input)
                name_without_ext = os.path.splitext(base_name)[0]
                libreoffice_output = os.path.join(abs_output_dir, f"{name_without_ext}.pdf")
                
                # Если имя отличается от требуемого, переименовываем
                if libreoffice_output != output_path and os.path.exists(libreoffice_output):
                    shutil.move(libreoffice_output, output_path)
                
                conversion_success = os.path.exists(output_path) and os.path.getsize(output_path) > 0
                if conversion_success:
                    logger.info("Конвертация с помощью LibreOffice успешна")
            else:
                error_message += f"LibreOffice: {process.stderr.decode('utf-8', errors='ignore')}; "
                logger.warning(f"LibreOffice вернул ошибку: {process.stderr.decode('utf-8', errors='ignore')}")
        except Exception as e:
            error_message += f"LibreOffice: {str(e)}; "
            logger.warning(f"Ошибка при использовании LibreOffice: {str(e)}")
    
    # Способ 3: Использование unoconv (через LibreOffice/OpenOffice)
    if UNOCONV_AVAILABLE and not conversion_success:
        try:
            logger.info(f"Попытка конвертации с помощью unoconv: {input_path}")
            unoconv.convert(input_path, output_path, 'pdf')
            conversion_success = os.path.exists(output_path) and os.path.getsize(output_path) > 0
            if conversion_success:
                logger.info("Конвертация с помощью unoconv успешна")
        except Exception as e:
            error_message += f"unoconv: {str(e)}; "
            logger.warning(f"unoconv не смог конвертировать файл: {str(e)}")
    
    if not conversion_success:
        logger.error(f"Все методы конвертации завершились неудачно: {error_message}")
    
    return conversion_success

# Настройка логирования
logger = logging.getLogger(__name__)

@login_required
def document_detail(request, document_id):
    document = get_object_or_404(Document, id=document_id)
    user = request.user
    organization = user.organization

    # Проверка доступа
    has_access = False
    if organization.is_prime_tech:
        has_access = (
            document.sender_organization == organization or
            document.recipient_organization == organization
        )
    else:
        has_access = (
            (document.sender == user.userprofile or document.recipient == user) and
            (document.sender_organization == organization or document.recipient_organization == organization)
        )

    if not has_access:
        django_messages.error(request, ("You do not have permission to view this document."))
        return redirect('staffs:dashboard')

    # Логирование просмотра документа
    UserActionLog.objects.create(
        user=user,
        action_type='view_document',
        details=f"Viewed document '{document.document_name}' (ID: {document.id})",
        performed_by=user
    )

    # Установка date_received, если пользователь — получатель и дата ещё не установлена
    if user == document.recipient and not document.date_received:
        document.date_received = timezone.now()
        document.save()

    page_images_base64 = []
    if document.document_content:
        file_path = document.document_content.path
        logger.info(f"Processing file: {file_path}")

        if not os.path.exists(file_path):
            logger.error(f"File not found on server: {file_path}")
            django_messages.error(request, ("File not found on server: ") + file_path)
            page_images_base64 = None
        else:
            content_type, _ = mimetypes.guess_type(file_path)
            logger.info(f"Detected content type: {content_type}")

            try:
                if content_type == 'application/pdf':
                    images = convert_from_path(file_path, dpi=200)
                    for img in images:
                        buffered = io.BytesIO()
                        img.save(buffered, format="PNG")
                        img_base64 = base64.b64encode(buffered.getvalue()).decode('utf-8')
                        page_images_base64.append(img_base64)
                elif content_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword']:
                    # Создаем временный файл для PDF
                    with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_pdf:
                        temp_pdf_path = temp_pdf.name
                    
                    try:
                        # Используем улучшенную функцию конвертации
                        conversion_success = convert_docx_to_pdf(file_path, temp_pdf_path)
                        
                        # Если конвертация удалась, генерируем изображения страниц
                        if conversion_success:
                            # Используем более высокое разрешение для лучшего качества
                            images = convert_from_path(temp_pdf_path, dpi=300)
                            for img in images:
                                buffered = io.BytesIO()
                                img.save(buffered, format="PNG", quality=95)  # Высокое качество
                                img_base64 = base64.b64encode(buffered.getvalue()).decode('utf-8')
                                page_images_base64.append(img_base64)
                        else:
                            # Если конвертация не удалась, используем резервный метод text_to_images
                            logger.warning(f"Failed to convert DOCX to PDF, using fallback method")
                            doc = DocxDocument(file_path)
                            text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
                            if not text:
                                logger.warning(f"DOCX file is empty: {file_path}")
                                django_messages.warning(request, ("The DOCX file is empty or contains no readable text."))
                            else:
                                images = text_to_images(text, width=800, height=1200)
                                for img in images:
                                    buffered = io.BytesIO()
                                    img.save(buffered, format="PNG")
                                    img_base64 = base64.b64encode(buffered.getvalue()).decode('utf-8')
                                    page_images_base64.append(img_base64)
                    finally:
                        # Удаляем временный PDF файл
                        if os.path.exists(temp_pdf_path):
                            os.unlink(temp_pdf_path)
                elif content_type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet':
                    df = pd.read_excel(file_path)
                    text = df.to_string(index=False)
                    images = text_to_images(text, width=800, height=1200)
                    for img in images:
                        buffered = io.BytesIO()
                        img.save(buffered, format="PNG")
                        img_base64 = base64.b64encode(buffered.getvalue()).decode('utf-8')
                        page_images_base64.append(img_base64)
                elif content_type == 'text/plain':
                    with open(file_path, 'r', encoding='utf-8') as f:
                        text = f.read()
                    images = text_to_images(text, width=800, height=1200)
                    for img in images:
                        buffered = io.BytesIO()
                        img.save(buffered, format="PNG")
                        img_base64 = base64.b64encode(buffered.getvalue()).decode('utf-8')
                        page_images_base64.append(img_base64)
                else:
                    logger.warning(f"Unsupported file type: {content_type} for file {file_path}")
                    django_messages.warning(request, _("Preview not available for this file type."))
                    page_images_base64 = None
            except Exception as e:
                logger.error(f"Error generating preview for {file_path}: {str(e)}", exc_info=True)
                page_images_base64 = None
                django_messages.error(request, _("Error generating preview: ") + str(e))
    else:
        logger.warning(f"No file attached to document ID {document_id}")
        django_messages.warning(request, ("No file attached to this document."))
        page_images_base64 = None

    if page_images_base64:
        page_data = [{'base64': img, 'download_url': f'/staffs/document/{document_id}/download-page/{i}/'} for i, img in enumerate(page_images_base64)]
    else:
        page_data = None

    context = {
        'document': document,
        'page_data': page_data,
        'is_prime_tech': organization.is_prime_tech,
        'status_choices': Document.STATUS_CHOICES,
    }
    return render(request, 'staffs/document_detail.html', context)

@login_required
def download_page(request, document_id, page_index):
    document = get_object_or_404(Document, id=document_id)
    user = request.user
    organization = user.organization

    has_access = False
    if organization.is_prime_tech:
        has_access = (
            document.sender_organization == organization or
            document.recipient_organization == organization
        )
    else:
        has_access = (
            (document.sender == user.userprofile or document.recipient == user) and
            (document.sender_organization == organization or document.recipient_organization == organization)
        )

    if not has_access:
        django_messages.error(request, ("You do not have permission to download this page."))
        return redirect('staffs:document-detail', document_id=document_id)

    file_path = document.document_content.path
    content_type, _ = mimetypes.guess_type(file_path)
    page_images = []

    try:
        if content_type == 'application/pdf':
            images = convert_from_path(file_path, dpi=200)
            page_images = images
        elif content_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword']:
            # Создаем временный файл для PDF
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_pdf:
                temp_pdf_path = temp_pdf.name
            
            try:
                # Используем улучшенную функцию конвертации
                conversion_success = convert_docx_to_pdf(file_path, temp_pdf_path)
                
                # Если конвертация удалась, используем изображения из PDF
                if conversion_success:
                    images = convert_from_path(temp_pdf_path, dpi=300)
                    page_images = images
                else:
                    # Если конвертация не удалась, используем резервный метод
                    doc = DocxDocument(file_path)
                    text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
                    images = text_to_images(text, width=800, height=1200)
                    page_images = images
            finally:
                # Удаляем временный PDF файл
                if os.path.exists(temp_pdf_path):
                    os.unlink(temp_pdf_path)
        elif content_type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet':
            df = pd.read_excel(file_path)
            text = df.to_string(index=False)
            images = text_to_images(text, width=800, height=1200)
            page_images = images
        elif content_type == 'text/plain':
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            images = text_to_images(text, width=800, height=1200)
            page_images = images

        if page_index < len(page_images):
            buffered = io.BytesIO()
            page_images[page_index].save(buffered, format="PNG")
            buffered.seek(0)
            return FileResponse(buffered, as_attachment=True, filename=f'page_{page_index}.png')
        else:
            django_messages.error(request, _("Page not found."))
            return redirect('staffs:document-detail', document_id=document_id)
    except Exception as e:
        django_messages.error(request, _("Error generating page for download: ") + str(e))
        return redirect('staffs:document-detail', document_id=document_id)


def text_to_images(text, width=800, height=1200):
    images = []
    try:
        # Используем Arial для лучшего отображения
        font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 14)
        bold_font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 14)
    except:
        font = ImageFont.load_default()
        bold_font = font

    # Разбиваем текст на параграфы
    paragraphs = text.split('\n')
    
    current_page_lines = []
    y = 40  # Начинаем с отступом сверху
    
    for paragraph in paragraphs:
        if not paragraph.strip():
            # Пустая строка - добавляем дополнительный отступ
            y += 20
            if y + 20 > height - 40:
                # Создаем новую страницу
                img = Image.new('RGB', (width, height), color='white')
                draw = ImageDraw.Draw(img)
                
                # Отрисовываем текущие линии
                y_pos = 40
                for line_info in current_page_lines:
                    line_text, is_bold = line_info
                    current_font = bold_font if is_bold else font
                    draw.text((40, y_pos), line_text, font=current_font, fill='black')
                    y_pos += 20
                
                images.append(img)
                current_page_lines = []
                y = 40
            continue

        # Определяем, является ли параграф заголовком (если начинается с # или содержит только заглавные буквы)
        is_heading = paragraph.startswith('#') or (paragraph.isupper() and len(paragraph) > 3)
        if is_heading:
            paragraph = paragraph.lstrip('#').strip()
            
        # Разбиваем длинные строки на части, которые помещаются по ширине
        words = paragraph.split()
        current_line = ""
        
        for word in words:
            test_line = current_line + " " + word if current_line else word
            # Используем соответствующий шрифт для измерения
            current_font = bold_font if is_heading else font
            if ImageDraw.Draw(Image.new('RGB', (1, 1))).textlength(test_line, font=current_font) < (width - 80):
                current_line = test_line
            else:
                # Добавляем заполненную строку
                if y + 20 > height - 40:
                    # Создаем новую страницу
                    img = Image.new('RGB', (width, height), color='white')
                    draw = ImageDraw.Draw(img)
                    
                    # Отрисовываем текущие линии
                    y_pos = 40
                    for line_info in current_page_lines:
                        line_text, is_bold = line_info
                        current_font = bold_font if is_bold else font
                        draw.text((40, y_pos), line_text, font=current_font, fill='black')
                        y_pos += 20
                    
                    images.append(img)
                    current_page_lines = []
                    y = 40
                
                current_page_lines.append((current_line, is_heading))
                y += 20
                current_line = word
        
        # Добавляем последнюю строку параграфа
        if current_line:
            if y + 20 > height - 40:
                img = Image.new('RGB', (width, height), color='white')
                draw = ImageDraw.Draw(img)
                y_pos = 40
                for line_info in current_page_lines:
                    line_text, is_bold = line_info
                    current_font = bold_font if is_bold else font
                    draw.text((40, y_pos), line_text, font=current_font, fill='black')
                    y_pos += 20
                images.append(img)
                current_page_lines = []
                y = 40
            
            current_page_lines.append((current_line, is_heading))
            y += 20
        
        # Добавляем отступ после параграфа
        y += 10

    # Создаем последнюю страницу, если остались непустые строки
    if current_page_lines:
        img = Image.new('RGB', (width, height), color='white')
        draw = ImageDraw.Draw(img)
        y_pos = 40
        for line_info in current_page_lines:
            line_text, is_bold = line_info
            current_font = bold_font if is_bold else font
            draw.text((40, y_pos), line_text, font=current_font, fill='black')
            y_pos += 20
        images.append(img)

    return images if images else [Image.new('RGB', (width, height), color='white')]


@login_required
def dashboard(request):
    user = request.user

    # Получаем активную вкладку из GET-параметра (по умолчанию 'sent')
    active_tab = request.GET.get('tab', 'sent')

    # Получаем отправленные и полученные документы
    sent_documents = Document.objects.filter(sender__user=user)
    received_documents = Document.objects.filter(recipient=user)

    # Статистика
    stats = {
        'total_sent': sent_documents.count(),
        'total_received': received_documents.count(),
        'total_draft': sent_documents.filter(status='draft').count() + received_documents.filter(status='draft').count(),
        'total_archived': sent_documents.filter(status='archived').count() + received_documents.filter(status='archived').count(),
        'total_documents': sent_documents.count() + received_documents.count(),
        'sent_status': sent_documents.filter(status='sent').count() + received_documents.filter(status='sent').count(),
        'received_status': sent_documents.filter(status='received').count() + received_documents.filter(status='received').count(),
    }

    # Фильтры
    status_filter = request.GET.get('status', '')
    start_date = request.GET.get('start_date', None)
    end_date = request.GET.get('end_date', None)
    org_filter = request.GET.get('org', '')
    sort_by = request.GET.get('sort_by', 'date_created')
    sort_order = request.GET.get('sort_order', 'desc')

    # Инициализация документов для активной вкладки
    documents = None
    page_obj = None

    if active_tab == 'sent':
        documents = sent_documents
    else:  # active_tab == 'received'
        documents = received_documents

    # Применяем фильтры
    if status_filter:
        documents = documents.filter(status=status_filter)
    if start_date:
        try:
            start_date = datetime.strptime(start_date, '%Y-%m-%d')
            documents = documents.filter(date_created__gte=start_date)
        except ValueError:
            django_messages.error(request, _("Invalid start date format. Use YYYY-MM-DD."))
    if end_date:
        try:
            end_date = datetime.strptime(end_date, '%Y-%m-%d')
            documents = documents.filter(date_created__lte=end_date)
        except ValueError:
            django_messages.error(request, _("Invalid end date format. Use YYYY-MM-DD."))
    if org_filter:
        documents = documents.filter(Q(sender_organization__id=org_filter) | Q(recipient_organization__id=org_filter))

    # Сортировка
    if sort_by == 'status':
        sort_field = 'status'
    else:
        sort_field = 'date_created'
    if sort_order == 'asc':
        documents = documents.order_by(sort_field)
    else:
        documents = documents.order_by(f'-{sort_field}')

    # Пагинация
    paginator = Paginator(documents, 9)  # 9 документов на страницу (3 ряда по 3 карточки)
    page_number = request.GET.get('page')
    try:
        page_obj = paginator.page(page_number)
    except PageNotAnInteger:
        page_obj = paginator.page(1)
    except EmptyPage:
        page_obj = paginator.page(paginator.num_pages)

    # Получаем список организаций для фильтра
    organizations = Organization.objects.all()

    # Список статусов для фильтра
    status_choices = Document.STATUS_CHOICES

    context = {
        'stats': stats,
        'documents': page_obj,
        'page_obj': page_obj,
        'status_choices': status_choices,
        'current_status': status_filter,
        'current_start_date': start_date,
        'current_end_date': end_date,
        'organizations': organizations,
        'current_org': org_filter,
        'sort_by': sort_by,
        'sort_order': sort_order,
        'active_tab': active_tab,
    }
    return render(request, 'staffs/dashboard.html', context)


@login_required
def send_document(request):
    if request.method == 'POST':
        form = SendDocumentForm(request.POST, request.FILES, user=request.user)
        if form.is_valid():
            document = form.save(commit=False)
            document.sender = request.user.userprofile
            document.sender_organization = request.user.organization
            
            # Автоматически устанавливаем организацию получателя
            recipient_user = document.recipient
            if recipient_user:
                document.recipient_organization = recipient_user.organization
            
            # Извлекаем информацию из документа
            if document.document_content:
                try:
                    # Определяем тип файла
                    file_path = document.document_content.path
                    content_type = mimetypes.guess_type(file_path)[0]

                    # Извлекаем текст и информацию в зависимости от типа файла
                    if content_type == 'application/pdf':
                        reader = PdfReader(file_path)
                        # Получаем количество страниц
                        document.page_count = len(reader.pages)
                        # Получаем текст из первой страницы для краткого описания
                        first_page = reader.pages[0]
                        text = first_page.extract_text()
                        # Берем первые 200 символов для краткого описания
                        document.summary = text[:200] if text else None

                    elif content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                        doc = DocxDocument(file_path)
                        # Получаем количество параграфов как страницы
                        document.page_count = len(doc.paragraphs)
                        # Получаем текст из первого параграфа для краткого описания
                        text = doc.paragraphs[0].text if doc.paragraphs else None
                        document.summary = text[:200] if text else None

                    elif content_type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet':
                        df = pd.read_excel(file_path)
                        # Считаем количество листов как страницы
                        wb = openpyxl.load_workbook(file_path)
                        document.page_count = len(wb.sheetnames)
                        # Получаем первые несколько строк для краткого описания
                        text = df.head().to_string()
                        document.summary = text[:200] if text else None

                    elif content_type == 'text/plain':
                        with open(file_path, 'r', encoding='utf-8') as f:
                            text = f.read()
                            # Считаем количество строк как страницы
                            document.page_count = len(text.splitlines())
                            # Берем первые 200 символов для краткого описания
                            document.summary = text[:200] if text else None

                    # Устанавливаем имя документа, если оно не задано
                    if not document.document_name:
                        document.document_name = os.path.basename(file_path)

                    # Устанавливаем метод отправки по умолчанию
                    if not document.method:
                        document.method = 'Внутренняя система'

                except Exception as e:
                    logger.error(f"Error extracting document info: {str(e)}")
                    # В случае ошибки устанавливаем базовые значения
                    if not document.summary:
                        document.summary = _("Document content could not be extracted")
                    if not document.page_count:
                        document.page_count = 1

            document.status = 'sent'
            document.date_sent = timezone.now()
            document.save()

            # Логирование действия отправки документа
            UserActionLog.objects.create(
                user=request.user,
                action_type='send_document',
                details=f"Sent document '{document.document_name}' to {document.recipient.username} ({document.recipient_organization.name if document.recipient_organization else 'No organization'})",
                performed_by=request.user
            )

            # Создаём уведомление для получателя
            if document.recipient:
                Notification.objects.create(
                    user=document.recipient,
                    message=f"New document '{document.document_name}' received from {document.sender.user.username} ({document.sender_organization.name})"
                )

            django_messages.success(request, _("Document sent successfully."))
            return redirect('staffs:dashboard')
        else:
            for field, errors in form.errors.items():
                for error in errors:
                    django_messages.error(request, f"{field}: {error}")
    else:
        form = SendDocumentForm(user=request.user)
    return render(request, 'staffs/send.html', {'form': form})


def landing_page(request):
    return render(request, 'landing_page.html')


def send_or_receive_view(request):
    return render(request, 'send_or_receive.html')


@login_required
def receive_document(request):
    documents = Document.objects.filter(recipient=request.user)
    updated = False
    for doc in documents:
        if doc.status == 'sent' and not doc.date_received:
            doc.status = 'received'
            doc.date_received = timezone.now()
            doc.save()
            updated = True

            # Логирование действия получения документа
            UserActionLog.objects.create(
                user=request.user,
                action_type='receive_document',
                details=f"Received document '{doc.document_name}' from {doc.sender.user.username}",
                performed_by=request.user
            )

    if updated:
        django_messages.success(request, _("Some documents have been marked as received."))

    # Фильтрация по статусу
    status_filter = request.GET.get('status', '')
    if status_filter:
        documents = documents.filter(status=status_filter)

    # Пагинация
    paginator = Paginator(documents, 6)
    page_number = request.GET.get('page')
    try:
        page_obj = paginator.page(page_number)
    except PageNotAnInteger:
        page_obj = paginator.page(1)
    except EmptyPage:
        page_obj = paginator.page(paginator.num_pages)

    context = {
        'documents': page_obj,
        'status_choices': Document.STATUS_CHOICES,
        'current_status': status_filter,
        'page_obj': page_obj,
    }
    return render(request, 'staffs/receive.html', context)


@login_required
def add_user(request):
    if request.user.role != 'admin':
        django_messages.error(request, _("You do not have permission to add users."))
        return redirect('staffs:dashboard')
    if request.method == 'POST':
        form = CustomUserCreationForm(request.POST)
        if form.is_valid():
            user = form.save()

            # Логирование добавления пользователя
            UserActionLog.objects.create(
                user=user,
                action_type='add_user',
                details=f"Added new user '{user.username}' with role '{user.role}' in organization '{user.organization.name}'",
                performed_by=request.user
            )

            django_messages.success(request, _("User added successfully!"))
            return redirect('staffs:dashboard')
    else:
        form = CustomUserCreationForm()
    return render(request, 'staffs/add_user.html', {'form': form})


@login_required
def document_log(request):
    user = request.user
    organization = user.organization

    if not organization.is_prime_tech:
        django_messages.error(request, _("Only PrimeTech organizations can access this page."))
        return redirect('staffs:dashboard')

    # Обработка скачивания таблицы
    if 'download' in request.GET:
        documents = Document.objects.filter(
            Q(sender_organization=organization) | Q(recipient_organization=organization)
        ).order_by('date_sent')

        data = []
        for idx, doc in enumerate(documents, start=1):
            # Получаем количество страниц из файла, если возможно
            page_count = doc.page_count
            if doc.document_content and hasattr(doc.document_content, 'path'):
                try:
                    if doc.document_content.path.endswith('.pdf'):
                        reader = PdfReader(doc.document_content.path)
                        page_count = len(reader.pages)
                    elif doc.document_content.path.endswith('.docx'):
                        doc_file = DocxDocument(doc.document_content.path)
                        page_count = len(doc_file.paragraphs)
                except Exception:
                    pass

            data.append({
                'Исходящий номер': idx,
                'Дата исходящего номера и дату принятия': doc.date_sent.strftime('%d.%m.%Y') if doc.date_sent else '-',
                'Адресат': doc.recipient_organization.name if doc.recipient_organization else '-',
                'Краткое содержание': doc.summary or '-',
                'Количество страниц': page_count or 0,
                'Приложение': doc.document_name if doc.document_content else (doc.attachment or '-'),
                'Исполнитель': doc.sender_organization.name if doc.sender_organization else '-',
                'Способ отправки': doc.method or 'Внутренняя система',
                'Дата отправки': doc.date_sent.strftime('%d.%m.%Y') if doc.date_sent else '-',
                'Дата исполнения': doc.date_received.strftime('%d.%m.%Y') if doc.date_received else '-',
                'Отметка о выполнении': doc.note or '-',
            })

        df = pd.DataFrame(data)

        output = BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Document Log', index=False)

        workbook = openpyxl.load_workbook(output)
        worksheet = workbook['Document Log']

        for col in worksheet.columns:
            column_letter = col[0].column_letter
            worksheet[f'{column_letter}1'].font = Font(bold=True)
            worksheet[f'{column_letter}1'].alignment = Alignment(horizontal='center', vertical='center')
            worksheet.column_dimensions[column_letter].width = 20

        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        for row in worksheet.rows:
            for cell in row:
                cell.border = thin_border
                cell.alignment = Alignment(horizontal='center', vertical='center')

        output.seek(0)
        output = BytesIO()
        workbook.save(output)
        output.seek(0)

        response = HttpResponse(
            output.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = 'attachment; filename="document_log.xlsx"'
        return response

    # Формирование данных для отображения таблицы
    documents = Document.objects.filter(
        Q(sender_organization=organization) | Q(recipient_organization=organization)
    ).order_by('date_sent')

    table_data = []
    for idx, doc in enumerate(documents, start=1):
        # Получаем количество страниц из файла
        page_count = doc.page_count
        if doc.document_content and hasattr(doc.document_content, 'path'):
            try:
                if doc.document_content.path.endswith('.pdf'):
                    reader = PdfReader(doc.document_content.path)
                    page_count = len(reader.pages)
                elif doc.document_content.path.endswith('.docx'):
                    doc_file = DocxDocument(doc.document_content.path)
                    page_count = len(doc_file.paragraphs)
            except Exception:
                pass

        # Определяем названия организаций
        sender_org_name = doc.sender_organization.name if doc.sender_organization else '-'
        recipient_org_name = doc.recipient_organization.name if doc.recipient_organization else '-'

        table_data.append({
            'id': doc.id,
            'number': idx,
            'date_sent_accepted': doc.date_sent.strftime('%d.%m.%Y') if doc.date_sent else '-',
            'recipient': recipient_org_name,
            'summary': doc.summary or '-',
            'page_count': page_count or 0,
            'attachment': doc.document_name if doc.document_content else (doc.attachment or '-'),
            'sender': sender_org_name,
            'method': doc.method or 'Внутренняя система',
            'date_sent': doc.date_sent.strftime('%d.%m.%Y') if doc.date_sent else '-',
            'date_received': doc.date_received.strftime('%d.%m.%Y') if doc.date_received else '-',
            'note': doc.note or '-',
        })

    context = {
        'table_data': table_data,
    }
    return render(request, 'staffs/document_log.html', context)


@require_GET
@login_required
def get_suggestions(request):
    """Возвращает подсказки для автозаполнения полей."""
    field = request.GET.get('field')
    query = request.GET.get('query', '').strip()
    
    if not query or len(query) < 2:
        return JsonResponse({'suggestions': []})

    suggestions = []
    if field == 'recipient':
        # Поиск по организациям
        orgs = Organization.objects.filter(
            name__icontains=query
        ).values_list('name', flat=True)[:10]
        suggestions.extend(list(orgs))

    elif field == 'sender':
        # Поиск по пользователям
        users = User.objects.filter(
            Q(first_name__icontains=query) |
            Q(last_name__icontains=query) |
            Q(username__icontains=query)
        )[:10]
        suggestions.extend([f"{user.get_full_name()} ({user.username})" for user in users])

    elif field == 'method':
        # Предопределенные методы отправки
        methods = ['Внутренняя система', 'Email', 'Почта России', 'Курьер', 'Факс']
        suggestions.extend([m for m in methods if query.lower() in m.lower()])

    return JsonResponse({'suggestions': suggestions})


@require_POST
@login_required
def update_document_field(request):
    document_id = request.POST.get('document_id')
    field = request.POST.get('field')
    value = request.POST.get('value')

    try:
        document = Document.objects.get(id=document_id)
        user = request.user
        organization = user.organization

        if not organization.is_prime_tech:
            return JsonResponse({'status': 'error', 'message': _("Only PrimeTech organizations can edit this table.")}, status=403)

        # Валидация и обновление поля
        if field == 'date_sent_accepted' or field == 'date_sent' or field == 'date_received':
            if value == '-':
                if field == 'date_sent_accepted' or field == 'date_sent':
                    document.date_sent = None
                elif field == 'date_received':
                    document.date_received = None
            else:
                try:
                    date_value = datetime.strptime(value, '%d.%m.%Y')
                    if field == 'date_sent_accepted' or field == 'date_sent':
                        document.date_sent = date_value
                    elif field == 'date_received':
                        document.date_received = date_value
                except ValueError:
                    return JsonResponse({'status': 'error', 'message': _("Invalid date format. Use DD.MM.YYYY.")}, status=400)
        elif field == 'page_count':
            try:
                page_count = int(value)
                if page_count < 1:
                    return JsonResponse({'status': 'error', 'message': _("Page count must be a positive integer.")}, status=400)
                document.page_count = page_count
            except ValueError:
                return JsonResponse({'status': 'error', 'message': _("Page count must be a number.")}, status=400)
        elif field == 'summary':
            document.summary = value
        elif field == 'method':
            document.method = value
        elif field == 'recipient_name':
            document.recipient_name = value
        elif field == 'attachment':
            document.attachment = value
        elif field == 'note':
            document.note = value

        document.save()
        return JsonResponse({'status': 'success', 'message': _("Field updated successfully.")})

    except Document.DoesNotExist:
        return JsonResponse({'status': 'error', 'message': _("Document not found.")}, status=404)
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)


@require_POST
@login_required
def change_document_status(request):
    if not request.user.organization.is_prime_tech:
        return JsonResponse({'status': 'error', 'message': _("Only PrimeTech organizations can change document status.")}, status=403)

    document_id = request.POST.get('document_id')
    new_status = request.POST.get('status')

    try:
        document = Document.objects.get(id=document_id)
        old_status = document.status
        if new_status in dict(Document.STATUS_CHOICES):
            document.status = new_status
            document.add_status_change_log(request.user, old_status, new_status)
            document.save()
            return JsonResponse({'status': 'success', 'message': _("Status updated successfully.")})
        return JsonResponse({'status': 'error', 'message': _("Invalid status.")}, status=400)
    except Document.DoesNotExist:
        return JsonResponse({'status': 'error', 'message': _("Document not found.")}, status=404)
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)


@login_required
def status_log_console(request):
    user = request.user
    organization = user.organization

    if not organization.is_prime_tech:
        django_messages.error(request, _("Only PrimeTech organizations can access this page."))
        return redirect('staffs:dashboard')

    # Проверка на запрос экспорта
    if 'export' in request.GET and request.GET['export'] == 'csv':
        documents = Document.objects.filter(
            Q(sender_organization=organization) | Q(recipient_organization=organization)
        )
        log_entries = []
        for doc in documents:
            if doc.status_change_log:
                entries = doc.status_change_log.strip().split('\n')
                for entry in entries:
                    if entry:
                        try:
                            timestamp_str, rest = entry.split(' - ', 1)
                            username, action = rest.split(' changed status from ', 1)
                            old_status, new_status = action.split(' to ')
                            timestamp = datetime.strptime(timestamp_str, '%Y-%m-%d %H:%M:%S')
                            log_entries.append({
                                'document': doc,
                                'timestamp': timestamp,
                                'username': username,
                                'old_status': old_status,
                                'new_status': new_status,
                            })
                        except (ValueError, IndexError):
                            continue

        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = 'attachment; filename="status_change_logs.csv"'
        writer = csv.writer(response)
        writer.writerow(['Timestamp', 'Document Name', 'User', 'Old Status', 'New Status'])
        for entry in log_entries:
            writer.writerow([
                entry['timestamp'],
                entry['document'].document_name,
                entry['username'],
                entry['old_status'],
                entry['new_status'],
            ])
        return response

    documents = Document.objects.filter(
        Q(sender_organization=organization) | Q(recipient_organization=organization)
    )

    start_date = request.GET.get('start_date', None)
    end_date = request.GET.get('end_date', None)
    user_filter = request.GET.get('user', None)
    document_filter = request.GET.get('document', None)

    log_entries = []
    for doc in documents:
        if doc.status_change_log:
            entries = doc.status_change_log.strip().split('\n')
            for entry in entries:
                if entry:
                    try:
                        timestamp_str, rest = entry.split(' - ', 1)
                        username, action = rest.split(' changed status from ', 1)
                        old_status, new_status = action.split(' to ')
                        timestamp = datetime.strptime(timestamp_str, '%Y-%m-%d %H:%M:%S')
                        log_entries.append({
                            'document': doc,
                            'timestamp': timestamp,
                            'username': username,
                            'old_status': old_status,
                            'new_status': new_status,
                        })
                    except (ValueError, IndexError):
                        continue

    # Фильтр по диапазону дат
    if start_date:
        try:
            start_date = datetime.strptime(start_date, '%Y-%m-%d')
            log_entries = [
                entry for entry in log_entries
                if entry['timestamp'].date() >= start_date.date()
            ]
        except ValueError:
            django_messages.error(request, _("Invalid start date format. Use YYYY-MM-DD."))

    if end_date:
        try:
            end_date = datetime.strptime(end_date, '%Y-%m-%d')
            log_entries = [
                entry for entry in log_entries
                if entry['timestamp'].date() <= end_date.date()
            ]
        except ValueError:
            django_messages.error(request, _("Invalid end date format. Use YYYY-MM-DD."))

    if user_filter:
        log_entries = [
            entry for entry in log_entries
            if user_filter.lower() in entry['username'].lower()
        ]

    if document_filter:
        log_entries = [
            entry for entry in log_entries
            if document_filter.lower() in entry['document'].document_name.lower()
        ]

    log_entries.sort(key=lambda x: x['timestamp'], reverse=True)

    paginator = Paginator(log_entries, 10)
    page_number = request.GET.get('page')
    try:
        page_obj = paginator.page(page_number)
    except PageNotAnInteger:
        page_obj = paginator.page(1)
    except EmptyPage:
        page_obj = paginator.page(paginator.num_pages)

    context = {
        'log_entries': page_obj,
        'page_obj': page_obj,
        'current_start_date': start_date,
        'current_end_date': end_date,
        'current_user': user_filter,
        'current_document': document_filter,
    }
    return render(request, 'staffs/status_log_console.html', context)


@login_required
def user_management(request):
    user = request.user

    if user.role != 'admin':
        django_messages.error(request, _("Only admins can access this page."))
        return redirect('staffs:dashboard')

    # Определяем активную вкладку
    active_tab = request.GET.get('tab', 'users')

    users = User.objects.all()
    organizations = Organization.objects.all()

    org_filter = request.GET.get('org', None)
    role_filter = request.GET.get('role', None)

    if org_filter:
        users = users.filter(organization__id=org_filter)

    if role_filter:
        users = users.filter(role=role_filter)

    paginator = Paginator(users, 10)
    page_number = request.GET.get('page')
    try:
        page_obj = paginator.page(page_number)
    except PageNotAnInteger:
        page_obj = paginator.page(1)
    except EmptyPage:
        page_obj = paginator.page(paginator.num_pages)

    # Статистика
    stats = {
        'total_users': User.objects.count(),
        'admins': User.objects.filter(role='admin').count(),
        'managers': User.objects.filter(role='manager').count(),
        'staff': User.objects.filter(role='staff').count(),
        'external': User.objects.filter(role='external').count(),
    }

    role_choices = User.ROLES

    context = {
        'users': page_obj,
        'page_obj': page_obj,
        'organizations': organizations,
        'role_choices': role_choices,
        'current_org': org_filter,
        'current_role': role_filter,
        'stats': stats,
        'active_tab': active_tab,
    }
    return render(request, 'staffs/user_management.html', context)


@require_POST
@login_required
def change_user_role(request):
    user = request.user
    if user.role != 'admin':
        return JsonResponse({'status': 'error', 'message': _("Only admins can change user roles.")}, status=403)

    user_id = request.POST.get('user_id')
    new_role = request.POST.get('role')

    try:
        target_user = User.objects.get(id=user_id)
        if new_role in dict(User.ROLES):
            old_role = target_user.role
            target_user.role = new_role
            target_user.save()
            # Логирование
            UserActionLog.objects.create(
                user=target_user,
                action_type='role_change',
                details=f"Role changed from {old_role} to {new_role}",
                performed_by=user
            )
            return JsonResponse({'status': 'success', 'message': _("User role updated successfully.")})
        return JsonResponse({'status': 'error', 'message': _("Invalid role.")}, status=400)
    except User.DoesNotExist:
        return JsonResponse({'status': 'error', 'message': _("User not found.")}, status=404)
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)


@require_POST
@login_required
def delete_user(request):
    user = request.user
    if user.role != 'admin':
        return JsonResponse({'status': 'error', 'message': _("Only admins can delete users.")}, status=403)

    user_id = request.POST.get('user_id')

    try:
        target_user = User.objects.get(id=user_id)
        if target_user == user:
            return JsonResponse({'status': 'error', 'message': _("You cannot delete yourself.")}, status=400)
        # Логирование перед удалением
        UserActionLog.objects.create(
            user=target_user,
            action_type='delete',
            details="User deleted",
            performed_by=user
        )
        target_user.delete()
        return JsonResponse({'status': 'success', 'message': _("User deleted successfully.")})
    except User.DoesNotExist:
        return JsonResponse({'status': 'error', 'message': _("User not found.")}, status=404)
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)


@login_required
def user_action_log(request):
    user = request.user

    if user.role != 'admin':
        django_messages.error(request, _("Only admins can access this page."))
        return redirect('staffs:dashboard')

    logs = UserActionLog.objects.all().order_by('-timestamp')

    # Фильтры
    user_filter = request.GET.get('user', None)
    action_filter = request.GET.get('action', None)

    if user_filter:
        logs = logs.filter(user__username__icontains=user_filter)

    if action_filter:
        logs = logs.filter(action_type=action_filter)

    # Пагинация
    paginator = Paginator(logs, 10)
    page_number = request.GET.get('page')
    try:
        page_obj = paginator.page(page_number)
    except PageNotAnInteger:
        page_obj = paginator.page(1)
    except EmptyPage:
        page_obj = paginator.page(paginator.num_pages)

    context = {
        'logs': page_obj,
        'page_obj': page_obj,
        'action_choices': UserActionLog.ACTION_TYPES,
        'current_user': user_filter,
        'current_action': action_filter,
    }
    return render(request, 'staffs/user_action_log.html', context)


@login_required
def notifications(request):
    # Получаем только непрочитанные уведомления
    notifications = request.user.notifications.filter(is_read=False).order_by('-created_at')
    
    if request.method == 'POST' and 'mark_as_read' in request.POST:
        notification_id = request.POST.get('notification_id')
        try:
            notification = request.user.notifications.filter(id=notification_id, is_read=False).first()
            if notification:
                notification.is_read = True
                notification.save()
                
                # Логируем действие
                UserActionLog.objects.create(
                    user=request.user,
                    action_type='read_notification',
                    details=f"Marked notification as read: {notification.message[:100]}...",
                    performed_by=request.user
                )
                
                return JsonResponse({
                    'status': 'success',
                    'message': _("Notification marked as read.")
                })
            return JsonResponse({
                'status': 'error',
                'message': _("Notification not found or already read.")
            }, status=404)
        except Exception as e:
            return JsonResponse({
                'status': 'error',
                'message': str(e)
            }, status=500)

    # Пагинация
    paginator = Paginator(notifications, 10)
    page_number = request.GET.get('page')
    try:
        page_obj = paginator.page(page_number)
    except PageNotAnInteger:
        page_obj = paginator.page(1)
    except EmptyPage:
        page_obj = paginator.page(paginator.num_pages)

    context = {
        'notifications': page_obj,
        'page_obj': page_obj,
        'unread_count': notifications.count(),  # Добавляем счетчик непрочитанных уведомлений
    }
    return render(request, 'staffs/notifications.html', context)


def is_prime_tech(user):
    return user.organization.is_prime_tech if user.is_authenticated and user.organization else False


@login_required
@user_passes_test(is_prime_tech)
def delete_document(request):
    if request.method == 'POST':
        document_id = request.POST.get('document_id')
        try:
            document = get_object_or_404(Document, id=document_id)

            # Проверка, что пользователь имеет доступ к этому документу (например, он из PrimeTech)
            if not request.user.organization.is_prime_tech:
                return JsonResponse({'status': 'error', 'message': 'Permission denied'}, status=403)

            # Логирование действия удаления
            UserActionLog.objects.create(
                user=request.user,
                action_type='delete_document',
                details=f"Deleted document '{document.document_name}' (ID: {document.id})",
                performed_by=request.user
            )

            # Удаление файла с диска, если он существует
            if document.document_content and os.path.isfile(document.document_content.path):
                os.remove(document.document_content.path)

            # Удаление записи документа из базы данных
            document.delete()

            return JsonResponse({'status': 'success', 'message': 'Document deleted successfully'})
        except Document.DoesNotExist:
            return JsonResponse({'status': 'error', 'message': 'Document not found'}, status=404)
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)}, status=500)
    return JsonResponse({'status': 'error', 'message': 'Invalid request method'}, status=400)


@login_required
def add_organization(request):
    if request.user.role != 'admin':
        django_messages.error(request, _("You do not have permission to add organizations."))
        return redirect('staffs:dashboard')

    if request.method == 'POST':
        form = OrganizationCreationForm(request.POST)
        if form.is_valid():
            organization = form.save()

            # Логирование добавления организации
            UserActionLog.objects.create(
                user=request.user,
                action_type='add_organization',
                details=f"Added new organization '{organization.name}'",
                performed_by=request.user
            )

            django_messages.success(request, _("Organization added successfully!"))
            return redirect(reverse('staffs:user_management') + '?tab=organizations')
    else:
        form = OrganizationCreationForm()

    return render(request, 'staffs/add_organization.html', {'form': form})


@require_GET
@login_required
def get_organization(request):
    if request.user.role != 'admin':
        return JsonResponse({'status': 'error', 'message': _("Only admins can access this data.")}, status=403)

    org_id = request.GET.get('org_id')
    try:
        organization = Organization.objects.get(id=org_id)
        return JsonResponse({
            'status': 'success',
            'organization': {
                'id': organization.id,
                'name': organization.name,
                'is_prime_tech': organization.is_prime_tech,
            }
        })
    except Organization.DoesNotExist:
        return JsonResponse({'status': 'error', 'message': _("Organization not found.")}, status=404)


@require_POST
@login_required
def edit_organization(request):
    if request.user.role != 'admin':
        return JsonResponse({'status': 'error', 'message': _("Only admins can edit organizations.")}, status=403)

    org_id = request.POST.get('org_id')
    try:
        organization = Organization.objects.get(id=org_id)
        form = OrganizationEditForm(request.POST, instance=organization)
        if form.is_valid():
            form.save()
            UserActionLog.objects.create(
                user=request.user,
                action_type='edit_organization',
                details=f"Edited organization '{organization.name}'",
                performed_by=request.user
            )
            return JsonResponse({'status': 'success', 'message': _("Organization updated successfully.")})
        else:
            errors = form.errors.as_json()
            return JsonResponse({'status': 'error', 'message': json.loads(errors)}, status=400)
    except Organization.DoesNotExist:
        return JsonResponse({'status': 'error', 'message': _("Organization not found.")}, status=404)
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)


@require_POST
@login_required
def delete_organization(request):
    if request.user.role != 'admin':
        return JsonResponse({'status': 'error', 'message': _("Only admins can delete organizations.")}, status=403)

    org_id = request.POST.get('org_id')
    try:
        organization = Organization.objects.get(id=org_id)
        if organization.users.exists():
            return JsonResponse({'status': 'error', 'message': _("Cannot delete organization with associated users.")}, status=400)
        if organization.is_prime_tech:
            return JsonResponse({'status': 'error', 'message': _("Cannot delete PrimeTech organization.")}, status=400)
        # Логирование перед удалением
        UserActionLog.objects.create(
            user=request.user,
            action_type='delete_organization',
            details=f"Deleted organization '{organization.name}'",
            performed_by=request.user
        )
        organization.delete()
        return JsonResponse({'status': 'success', 'message': _("Organization deleted successfully.")})
    except Organization.DoesNotExist:
        return JsonResponse({'status': 'error', 'message': _("Organization not found.")}, status=404)
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)


@require_GET
@login_required
def get_org_users(request):
    if request.user.role != 'admin':
        return JsonResponse({'status': 'error', 'message': _("Only admins can access this data.")}, status=403)

    org_id = request.GET.get('org_id')
    try:
        organization = Organization.objects.get(id=org_id)
        users = organization.users.all()
        users_data = [{
            'id': user.id,
            'username': user.username,
            'role': dict(User.ROLES).get(user.role, user.role),
        } for user in users]
        return JsonResponse({
            'status': 'success',
            'users': users_data
        })
    except Organization.DoesNotExist:
        return JsonResponse({'status': 'error', 'message': _("Organization not found.")}, status=404)


@login_required
def get_chats(request):
    if not request.user.is_authenticated:
        return JsonResponse({'status': 'error', 'message': 'User not authenticated'}, status=401)

    user_org = request.user.organization
    if user_org.is_prime_tech:

        chats = Chat.objects.filter(prime_tech_organization=user_org, is_support=False)
    else:
        chats = Chat.objects.filter(secondary_organization=user_org, is_support=False)

    chat_list = [{
        'id': chat.id,
        'name': chat.secondary_organization.name if chat.secondary_organization else chat.name,
        'last_message': chat.messages.last().message if chat.messages.exists() else 'No messages'
    } for chat in chats]
    return JsonResponse({'status': 'success', 'chats': chat_list})


@require_GET
def get_support_chat(request):
    try:
        chat = Chat.objects.get(is_support=True)
        session_key = request.session.session_key
        if not session_key:
            request.session.save()
            session_key = request.session.session_key
        return JsonResponse({
            'status': 'success',
            'chat': {
                'id': chat.id,
                'name': chat.name,
                'last_message': chat.messages.last().message if chat.messages.exists() else None
            }
        })
    except Chat.DoesNotExist:
        prime_tech_org = Organization.objects.filter(is_prime_tech=True).first()
        if prime_tech_org:
            chat = Chat.objects.create(
                prime_tech_organization=prime_tech_org,
                is_support=True,
                name="Support Chat"
            )
            return JsonResponse({
                'status': 'success',
                'chat': {
                    'id': chat.id,
                    'name': chat.name,
                    'last_message': None
                }
            })
        return JsonResponse({'status': 'error', 'message': _("Support chat not found.")}, status=404)


@login_required
def chat_history(request, chat_id):
    user = request.user
    organization = user.organization

    try:
        chat = Chat.objects.get(id=chat_id, is_support=False)
        if not organization:
            return JsonResponse({'status': 'error', 'message': _("User has no organization.")}, status=403)
        if organization.is_prime_tech:
            if chat.prime_tech_organization != organization:
                return JsonResponse({'status': 'error', 'message': _("You do not have access to this chat.")}, status=403)
        else:
            if chat.secondary_organization != organization:
                return JsonResponse({'status': 'error', 'message': _("You do not have access to this chat.")}, status=403)

        messages = chat.messages.all()
        messages_data = [{
            'sender': msg.sender.username if msg.sender else "Guest",
            'message': msg.message,
            'timestamp': msg.timestamp.strftime('%Y-%m-%d %H:%M:%S')
        } for msg in messages]

        return JsonResponse({'status': 'success', 'messages': messages_data})

    except Chat.DoesNotExist:
        return JsonResponse({'status': 'error', 'message': _("Chat not found.")}, status=404)


@require_GET
def support_chat_history(request):
    try:
        chat = Chat.objects.get(is_support=True)
        session_key = request.session.session_key
        if not session_key:
            request.session.save()
            session_key = request.session.session_key

        if request.user.is_authenticated and request.user.organization.is_prime_tech:
            messages = chat.messages.all()
        else:
            messages = chat.messages.filter(session_key=session_key) | chat.messages.filter(sender__organization__is_prime_tech=True)

        messages_data = [{
            'sender': msg.sender.username if msg.sender else "Guest",
            'message': msg.message,
            'timestamp': msg.timestamp.strftime('%Y-%m-%d %H:%M:%S')
        } for msg in messages.order_by('timestamp')]

        return JsonResponse({
            'status': 'success',
            'messages': messages_data
        })
    except Chat.DoesNotExist:
        return JsonResponse({'status': 'error', 'message': _("Support chat not found.")}, status=404)



@login_required
def backup_management(request):
    if request.user.role != 'admin':
        django_messages.error(request, _("Only admins can access this page."))
        return redirect('staffs:dashboard')

    # Определяем тип базы данных
    db_engine = settings.DATABASES['default']['ENGINE']
    
    # Путь к папке для резервных копий
    backup_dir = os.path.join(settings.MEDIA_ROOT, 'backups')
    if not os.path.exists(backup_dir):
        os.makedirs(backup_dir)

    # Получаем список резервных копий
    backups = []
    for filename in os.listdir(backup_dir):
        file_path = os.path.join(backup_dir, filename)
        if os.path.isfile(file_path) and filename.endswith('.sql'):
            created_at = datetime.fromtimestamp(os.path.getctime(file_path))
            size = os.path.getsize(file_path) / (1024 * 1024)  # Размер в MB
            
            # Определяем тип резервной копии с помощью вспомогательной функции
            backup_type = determine_sql_file_type(file_path)
                
            # Упрощенная логика совместимости: 
            # 1. Считаем все PostgreSQL резервные копии совместимыми с PostgreSQL БД
            # 2. Считаем все SQLite резервные копии совместимыми с SQLite БД
            # 3. Generic SQL и Unknown считаем совместимыми со всеми
            compatible = (backup_type in ["SQL (Generic)", "Unknown"]) or \
                   (('postgresql' in db_engine and 'PostgreSQL' in backup_type) or \
                    ('sqlite3' in db_engine and backup_type == 'SQLite'))
                
            backups.append({
                'filename': filename,
                'created_at': created_at,
                'size': size,
                'type': backup_type,
                'compatible': compatible
            })

    # Сортировка по дате создания (от новых к старым)
    backups.sort(key=lambda x: x['created_at'], reverse=True)

    # Статистика
    stats = {
        'total_backups': len(backups),
        'last_backup': backups[0]['created_at'].strftime('%Y-%m-%d %H:%M:%S') if backups else None,
        'storage_used': sum(b['size'] for b in backups),
        'current_db_type': 'SQLite' if 'sqlite3' in db_engine else ('PostgreSQL' if 'postgresql' in db_engine else 'Other'),
    }

    # Пагинация
    paginator = Paginator(backups, 9)  # 9 резервных копий на страницу
    page_number = request.GET.get('page')
    try:
        page_obj = paginator.page(page_number)
    except PageNotAnInteger:
        page_obj = paginator.page(1)
    except EmptyPage:
        page_obj = paginator.page(paginator.num_pages)

    context = {
        'backups': page_obj,
        'page_obj': page_obj,
        'stats': stats,
    }
    return render(request, 'staffs/backup.html', context)

@login_required
def create_backup(request):
    if request.user.role != 'admin':
        return JsonResponse({'status': 'error', 'message': _("Only admins can create backups.")}, status=403)

    if request.method == 'POST':
        try:
            # Определяем тип базы данных
            db_engine = settings.DATABASES['default']['ENGINE']
            db_name = settings.DATABASES['default']['NAME']
            
            print(f"Database engine: {db_engine}")
            print(f"Database name: {db_name}")
            
            # Создаем директорию для бэкапов, если ее нет
            backup_dir = os.path.join(settings.MEDIA_ROOT, 'backups')
            print(f"Backup dir: {backup_dir}")
            if not os.path.exists(backup_dir):
                os.makedirs(backup_dir)
                
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            backup_filename = f"backup_{timestamp}.sql"
            backup_path = os.path.join(backup_dir, backup_filename)
            print(f"Creating SQL backup: {backup_path}")
            
            # SQLite
            if 'sqlite3' in db_engine:
                print("Using SQLite backup method")
                if not os.path.exists(db_name):
                    return JsonResponse({'status': 'error', 'message': _("Database file not found.")}, status=404)
                
                # Создаем дамп SQLite
                with open(backup_path, 'w', encoding='utf-8') as f:
                    process = subprocess.run(['sqlite3', db_name, '.dump'], stdout=subprocess.PIPE, text=True, check=True)
                    f.write(process.stdout)
            
            # PostgreSQL
            elif 'postgresql' in db_engine:
                print("Using PostgreSQL backup method")
                db_user = settings.DATABASES['default']['USER']
                db_host = settings.DATABASES['default']['HOST']
                db_port = settings.DATABASES['default']['PORT']
                db_password = settings.DATABASES['default']['PASSWORD']
                
                # Устанавливаем переменную окружения для пароля PostgreSQL
                env = os.environ.copy()
                env['PGPASSWORD'] = db_password
                
                # Ищем pg_dump нужной версии
                pg_dump_paths = [
                    'pg_dump',                                        # Стандартный путь
                    '/opt/homebrew/opt/postgresql@15/bin/pg_dump',    # Homebrew PostgreSQL 15
                    '/opt/homebrew/opt/postgresql/bin/pg_dump',       # Текущая версия Homebrew
                    '/usr/local/bin/pg_dump',                         # Обычный путь для Homebrew
                    '/Applications/Postgres.app/Contents/Versions/15/bin/pg_dump', # Postgres.app
                ]
                
                pg_dump_cmd = None
                for path in pg_dump_paths:
                    try:
                        # Проверяем версию pg_dump
                        version_process = subprocess.run(
                            [path, '--version'], 
                            stdout=subprocess.PIPE, 
                            stderr=subprocess.PIPE, 
                            text=True,
                            env=env)
                        
                        if version_process.returncode == 0 and '15.' in version_process.stdout:
                            pg_dump_cmd = path
                            print(f"Using pg_dump: {pg_dump_cmd}")  # Отладка
                            break
                    except FileNotFoundError:
                        continue
                
                if not pg_dump_cmd:
                    pg_dump_cmd = 'pg_dump'  # Используем стандартный путь, если ничего не найдено
                    print("Warning: Could not find pg_dump 15.x, using default")
                
                # Выполняем команду pg_dump для создания резервной копии
                process = subprocess.run([
                    pg_dump_cmd,
                    '-h', db_host,
                    '-p', db_port,
                    '-U', db_user,
                    '-F', 'p',  # plain text format
                    '-f', backup_path,
                    db_name
                ], env=env, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
                
                # Проверяем, если команда вернула ошибку
                if process.returncode != 0:
                    print(f"pg_dump error: {process.stderr}")  # Отладка
                    return JsonResponse({'status': 'error', 'message': f"Failed to create backup: {process.stderr}"}, status=500)
            
            else:
                return JsonResponse({'status': 'error', 'message': f"Unsupported database engine: {db_engine}"}, status=500)
            
            print(f"Backup created successfully: {backup_filename}")  # Отладка
            UserActionLog.objects.create(
                user=request.user,
                action_type='create_backup',
                details=f"Created SQL backup '{backup_filename}'",
                performed_by=request.user
            )
            return JsonResponse({'status': 'success', 'message': _("Backup created successfully.")})
        except subprocess.CalledProcessError as e:
            print(f"Subprocess error: {str(e)}")  # Отладка
            return JsonResponse({'status': 'error', 'message': f"Failed to create SQL dump: {str(e)}"}, status=500)
        except Exception as e:
            print(f"General error: {str(e)}")  # Отладка
            return JsonResponse({'status': 'error', 'message': str(e)}, status=500)
    return JsonResponse({'status': 'error', 'message': _("Invalid request method.")}, status=400)

@login_required
def download_backup(request, filename):
    if request.user.role != 'admin':
        django_messages.error(request, _("Only admins can download backups."))
        return redirect('staffs:dashboard')

    backup_path = os.path.join(settings.MEDIA_ROOT, 'backups', filename)
    if not os.path.exists(backup_path):
        django_messages.error(request, _("Backup file not found."))
        return redirect('staffs:backup_management')

    # Логирование
    UserActionLog.objects.create(
        user=request.user,
        action_type='download_backup',
        details=f"Downloaded SQL backup '{filename}'",
        performed_by=request.user
    )

    response = FileResponse(open(backup_path, 'rb'), as_attachment=True, filename=filename)
    response['Content-Type'] = 'application/sql'
    return response

@login_required
def restore_backup(request):
    if request.user.role != 'admin':
        return JsonResponse({'status': 'error', 'message': _("Only admins can restore backups.")}, status=403)

    if request.method == 'POST':
        filename = request.POST.get('filename')
        if not filename:
            return JsonResponse({'status': 'error', 'message': _("No backup filename provided.")}, status=400)

        try:
            # Определяем тип базы данных
            db_engine = settings.DATABASES['default']['ENGINE']
            db_name = settings.DATABASES['default']['NAME']
            
            print(f"Database engine: {db_engine}")
            print(f"Database name: {db_name}")
            
            # Путь к файлу резервной копии
            backup_path = os.path.join(settings.MEDIA_ROOT, 'backups', filename)
            print(f"Restoring from backup: {backup_path}")
            if not os.path.exists(backup_path):
                return JsonResponse({'status': 'error', 'message': _("Backup file not found.")}, status=404)
            
            # Проверяем совместимость файла с текущей базой данных
            backup_type = determine_sql_file_type(backup_path)
            print(f"Determined backup type: {backup_type}")
            
            # Упрощенная логика совместимости: 
            # 1. Считаем все PostgreSQL резервные копии совместимыми с PostgreSQL БД
            # 2. Считаем все SQLite резервные копии совместимыми с SQLite БД
            # 3. Generic SQL и Unknown считаем совместимыми со всеми
            is_compatible = (backup_type in ["SQL (Generic)", "Unknown"]) or \
                            (('postgresql' in db_engine and 'PostgreSQL' in backup_type) or \
                             ('sqlite3' in db_engine and backup_type == 'SQLite'))
            
            if not is_compatible:
                print(f"Warning: Backup type '{backup_type}' may not be compatible with current database type '{db_engine}'")
                # Логируем предупреждение, но продолжаем восстановление

            # Создаем предварительную резервную копию перед восстановлением
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            pre_restore_backup = os.path.join(settings.MEDIA_ROOT, 'backups', f"pre_restore_{timestamp}.sql")
            
            # SQLite
            if 'sqlite3' in db_engine:
                print("Using SQLite restore method")
                
                # Проверяем права доступа к папке backups
                backups_dir = os.path.join(settings.MEDIA_ROOT, 'backups')
                if not os.path.exists(backups_dir):
                    print(f"Creating backups directory: {backups_dir}")
                    os.makedirs(backups_dir, exist_ok=True)

                if not os.access(backups_dir, os.W_OK):
                    print(f"No write permission for backups directory: {backups_dir}")
                    return JsonResponse({'status': 'error', 'message': _("Backups directory lacks write permissions.")}, status=500)

                # Проверяем права доступа к файлу базы данных
                if not os.access(db_name, os.W_OK):
                    print(f"No write permission for database: {db_name}")
                    return JsonResponse({'status': 'error', 'message': _("Database file is read-only or lacks write permissions.")}, status=500)

                # Проверяем права доступа к родительской папке базы данных
                db_dir = os.path.dirname(db_name)
                if not os.access(db_dir, os.W_OK):
                    print(f"No write permission for database directory: {db_dir}")
                    return JsonResponse({'status': 'error', 'message': _("Database directory lacks write permissions.")}, status=500)

                # Создаём резервную копию текущей базы перед восстановлением
                print(f"Creating pre-restore backup: {pre_restore_backup}")
                with open(pre_restore_backup, 'w', encoding='utf-8') as f:
                    process = subprocess.run(['sqlite3', db_name, '.dump'], stdout=subprocess.PIPE, text=True, check=True)
                    f.write(process.stdout)

                # Закрываем все соединения с базой данных
                print("Closing database connections")
                connections.close_all()

                # Создаём временную базу данных
                temp_db_path = os.path.join(settings.MEDIA_ROOT, f'temp_restore_{timestamp}.sqlite3')
                print(f"Creating temporary database: {temp_db_path}")

                # Проверяем права доступа к MEDIA_ROOT для временной базы
                if not os.access(settings.MEDIA_ROOT, os.W_OK):
                    print(f"No write permission for media directory: {settings.MEDIA_ROOT}")
                    return JsonResponse({'status': 'error', 'message': _("Media directory lacks write permissions.")}, status=500)

                # Создаём пустую временную базу
                open(temp_db_path, 'a').close()
                os.chmod(temp_db_path, 0o664)  # Устанавливаем права для временного файла

                # Восстанавливаем базу из SQL-дампа
                print(f"Restoring SQL dump to temporary database: {temp_db_path}")
                with open(backup_path, 'r', encoding='utf-8') as f:
                    sql_dump = f.read()
                    conn = sqlite3.connect(temp_db_path)
                    try:
                        conn.executescript(sql_dump)
                        conn.commit()
                    finally:
                        conn.close()

                # Заменяем текущую базу восстановленной
                print(f"Replacing current database with restored: {db_name}")
                os.replace(temp_db_path, db_name)
                
            # PostgreSQL
            elif 'postgresql' in db_engine:
                print("Using PostgreSQL restore method")
                db_user = settings.DATABASES['default']['USER']
                db_host = settings.DATABASES['default']['HOST']
                db_port = settings.DATABASES['default']['PORT']
                db_password = settings.DATABASES['default']['PASSWORD']
                
                # Устанавливаем переменную окружения для пароля PostgreSQL
                env = os.environ.copy()
                env['PGPASSWORD'] = db_password
                
                # Ищем pg_dump и psql нужной версии
                pg_dump_paths = [
                    'pg_dump',                                        # Стандартный путь
                    '/opt/homebrew/opt/postgresql@15/bin/pg_dump',    # Homebrew PostgreSQL 15
                    '/opt/homebrew/opt/postgresql/bin/pg_dump',       # Текущая версия Homebrew
                    '/usr/local/bin/pg_dump',                         # Обычный путь для Homebrew
                    '/Applications/Postgres.app/Contents/Versions/15/bin/pg_dump', # Postgres.app
                ]
                
                psql_paths = [
                    'psql',                                        # Стандартный путь
                    '/opt/homebrew/opt/postgresql@15/bin/psql',    # Homebrew PostgreSQL 15
                    '/opt/homebrew/opt/postgresql/bin/psql',       # Текущая версия Homebrew
                    '/usr/local/bin/psql',                         # Обычный путь для Homebrew
                    '/Applications/Postgres.app/Contents/Versions/15/bin/psql', # Postgres.app
                ]
                
                # Находим правильную версию pg_dump
                pg_dump_cmd = None
                for path in pg_dump_paths:
                    try:
                        version_process = subprocess.run(
                            [path, '--version'], 
                            stdout=subprocess.PIPE, 
                            stderr=subprocess.PIPE, 
                            text=True,
                            env=env)
                        
                        if version_process.returncode == 0 and '15.' in version_process.stdout:
                            pg_dump_cmd = path
                            print(f"Using pg_dump: {pg_dump_cmd}")  # Отладка
                            break
                    except FileNotFoundError:
                        continue
                
                if not pg_dump_cmd:
                    pg_dump_cmd = 'pg_dump'  # Используем стандартный путь, если ничего не найдено
                    print("Warning: Could not find pg_dump 15.x, using default")
                    
                # Находим правильную версию psql
                psql_cmd = None
                for path in psql_paths:
                    try:
                        version_process = subprocess.run(
                            [path, '--version'], 
                            stdout=subprocess.PIPE, 
                            stderr=subprocess.PIPE, 
                            text=True,
                            env=env)
                        
                        if version_process.returncode == 0 and '15.' in version_process.stdout:
                            psql_cmd = path
                            print(f"Using psql: {psql_cmd}")  # Отладка
                            break
                    except FileNotFoundError:
                        continue
                
                if not psql_cmd:
                    psql_cmd = 'psql'  # Используем стандартный путь, если ничего не найдено
                    print("Warning: Could not find psql 15.x, using default")
                
                # Создаем резервную копию текущего состояния перед восстановлением
                print(f"Creating pre-restore backup: {pre_restore_backup}")  # Отладка
                pre_backup_process = subprocess.run([
                    pg_dump_cmd,
                    '-h', db_host,
                    '-p', db_port,
                    '-U', db_user,
                    '-F', 'p',  # plain text format
                    '-f', pre_restore_backup,
                    db_name
                ], env=env, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
                
                if pre_backup_process.returncode != 0:
                    print(f"Pre-restore backup error: {pre_backup_process.stderr}")  # Отладка
                    return JsonResponse({'status': 'error', 'message': f"Failed to create pre-restore backup: {pre_backup_process.stderr}"}, status=500)
                
                # Закрываем все соединения с базой данных
                print("Closing database connections")  # Отладка
                connections.close_all()
                
                # Подготавливаем файл с опциями для psql
                options_file = os.path.join(settings.MEDIA_ROOT, f"psql_options_{timestamp}.txt")
                with open(options_file, 'w') as f:
                    f.write("SET session_replication_role = 'replica';\n")  # Отключаем триггеры и ограничения
                
                # Очищаем существующие данные перед восстановлением
                print("Cleaning database before restore")  # Отладка
                clean_process = subprocess.run([
                    psql_cmd,
                    '-h', db_host,
                    '-p', db_port,
                    '-U', db_user,
                    '-d', db_name,
                    '-c', "SELECT pg_terminate_backend(pid) FROM pg_stat_activity WHERE datname = current_database() AND pid <> pg_backend_pid(); DROP SCHEMA public CASCADE; CREATE SCHEMA public; GRANT ALL ON SCHEMA public TO postgres; GRANT ALL ON SCHEMA public TO public;"
                ], env=env, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
                
                if clean_process.returncode != 0:
                    os.unlink(options_file)
                    print(f"Error cleaning database: {clean_process.stderr}")  # Отладка
                    return JsonResponse({'status': 'error', 'message': f"Failed to clean database before restore: {clean_process.stderr}"}, status=500)
                
                # Восстанавливаем базу из SQL-дампа с дополнительными опциями
                print(f"Restoring database from backup")  # Отладка
                restore_process = subprocess.run([
                    psql_cmd,
                    '-h', db_host,
                    '-p', db_port,
                    '-U', db_user,
                    '-d', db_name,
                    '-f', options_file,
                    '-f', backup_path
                ], env=env, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
                
                # Удаляем временный файл с опциями
                os.unlink(options_file)
                
                if restore_process.returncode != 0:
                    print(f"Restore error: {restore_process.stderr}")  # Отладка
                    return JsonResponse({'status': 'error', 'message': f"Failed to restore database: {restore_process.stderr}"}, status=500)
                
                # Включаем триггеры и ограничения обратно
                post_restore_process = subprocess.run([
                    psql_cmd,
                    '-h', db_host,
                    '-p', db_port,
                    '-U', db_user,
                    '-d', db_name,
                    '-c', "SET session_replication_role = 'origin';"
                ], env=env, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
            
            else:
                return JsonResponse({'status': 'error', 'message': f"Unsupported database engine: {db_engine}"}, status=500)
            
            # Логирование
            UserActionLog.objects.create(
                user=request.user,
                action_type='restore_backup',
                details=f"Restored database from SQL backup '{filename}'",
                performed_by=request.user
            )

            print("Database restored successfully")  # Отладка
            return JsonResponse({'status': 'success', 'message': _("Database restored successfully.")})
        except subprocess.CalledProcessError as e:
            print(f"Subprocess error during restore: {str(e)}")  # Отладка
            return JsonResponse({'status': 'error', 'message': f"Failed to process SQL dump: {str(e)}"}, status=500)
        except sqlite3.Error as e:
            print(f"SQLite error during restore: {str(e)}")  # Отладка
            return JsonResponse({'status': 'error', 'message': f"SQLite error: {str(e)}"}, status=500)
        except Exception as e:
            print(f"General error during restore: {str(e)}")  # Отладка
            return JsonResponse({'status': 'error', 'message': str(e)}, status=500)
    return JsonResponse({'status': 'error', 'message': _("Invalid request method.")}, status=400)

@login_required
def delete_backup(request):
    if request.user.role != 'admin':
        return JsonResponse({'status': 'error', 'message': _("Only admins can delete backups.")}, status=403)

    if request.method == 'POST':
        filename = request.POST.get('filename')
        try:
            backup_path = os.path.join(settings.MEDIA_ROOT, 'backups', filename)
            if not os.path.exists(backup_path):
                return JsonResponse({'status': 'error', 'message': _("Backup file not found.")}, status=404)

            os.remove(backup_path)

            # Логирование
            UserActionLog.objects.create(
                user=request.user,
                action_type='delete_backup',
                details=f"Deleted SQL backup '{filename}'",
                performed_by=request.user
            )

            return JsonResponse({'status': 'success', 'message': _("Backup deleted successfully.")})
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)}, status=500)
    return JsonResponse({'status': 'error', 'message': _("Invalid request method.")}, status=400)
@require_POST
@login_required
def reset_document_data(request):
    """Сбрасывает данные документа к исходному состоянию."""
    if not request.user.organization.is_prime_tech:
        return JsonResponse({'status': 'error', 'message': _("Only PrimeTech organizations can reset document data.")}, status=403)

    try:
        # Получаем все документы организации
        documents = Document.objects.filter(
            Q(sender_organization=request.user.organization) | 
            Q(recipient_organization=request.user.organization)
        ).order_by('date_sent')

        table_data = []
        for idx, doc in enumerate(documents, start=1):
            # Получаем количество страниц из файла
            page_count = doc.page_count
            if doc.document_content and hasattr(doc.document_content, 'path'):
                try:
                    if doc.document_content.path.endswith('.pdf'):
                        from PyPDF2 import PdfReader
                        reader = PdfReader(doc.document_content.path)
                        page_count = len(reader.pages)
                    elif doc.document_content.path.endswith('.docx'):
                        from docx import Document as DocxDocument
                        doc_file = DocxDocument(doc.document_content.path)
                        page_count = len(doc_file.paragraphs)
                except Exception:
                    pass

            # Определяем названия организаций
            sender_org_name = doc.sender_organization.name if doc.sender_organization else '-'
            recipient_org_name = doc.recipient_organization.name if doc.recipient_organization else '-'

            # Формируем данные для таблицы
            table_data.append({
                'id': doc.id,
                'number': idx,
                'date_sent_accepted': doc.date_sent.strftime('%d.%m.%Y') if doc.date_sent else '-',
                'recipient': recipient_org_name,  # Используем название организации-получателя
                'summary': doc.summary or '-',
                'page_count': page_count or 0,
                'attachment': doc.document_name if doc.document_content else (doc.attachment or '-'),
                'sender': sender_org_name,  # Используем название организации-отправителя
                'method': doc.method or 'Внутренняя система',
                'date_sent': doc.date_sent.strftime('%d.%m.%Y') if doc.date_sent else '-',
                'date_received': doc.date_received.strftime('%d.%m.%Y') if doc.date_received else '-',
                'note': doc.note or '-',
            })

        # Логируем действие
        UserActionLog.objects.create(
            user=request.user,
            action_type='reset_document_data',
            details="Reset document data to original state",
            performed_by=request.user
        )

        return JsonResponse({
            'status': 'success',
            'message': _("Document data has been reset successfully."),
            'table_data': table_data
        })

    except Exception as e:
        return JsonResponse({
            'status': 'error',
            'message': str(e)
        }, status=500)

@require_POST
@login_required
def edit_user(request):
    if request.user.role != 'admin':
        return JsonResponse({'status': 'error', 'message': _("Only admins can edit users.")}, status=403)

    user_id = request.POST.get('user_id')
    try:
        user = User.objects.get(id=user_id)
        new_username = request.POST.get('username')
        new_role = request.POST.get('role')
        new_org_id = request.POST.get('organization')

        # Проверяем, не существует ли уже пользователь с таким именем
        if User.objects.filter(username=new_username).exclude(id=user_id).exists():
            return JsonResponse({
                'status': 'error',
                'message': _("A user with this username already exists.")
            }, status=400)

        # Обновляем данные пользователя
        user.username = new_username
        user.role = new_role
        
        # Обновляем организацию
        if new_org_id:
            try:
                organization = Organization.objects.get(id=new_org_id)
                user.organization = organization
            except Organization.DoesNotExist:
                return JsonResponse({
                    'status': 'error',
                    'message': _("Selected organization does not exist.")
                }, status=400)
        else:
            user.organization = None

        user.save()

        # Логируем изменение
        UserActionLog.objects.create(
            user=request.user,
            action_type='edit_user',
            details=f"Edited user '{user.username}'",
            performed_by=request.user
        )

        return JsonResponse({
            'status': 'success',
            'message': _("User updated successfully."),
            'user': {
                'username': user.username,
                'role': user.get_role_display(),
                'organization': user.organization.name if user.organization else None
            }
        })
    except User.DoesNotExist:
        return JsonResponse({'status': 'error', 'message': _("User not found.")}, status=404)
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@require_POST
@login_required
def change_password(request):
    if request.user.role != 'admin':
        return JsonResponse({'status': 'error', 'message': _("Only admins can change user passwords.")}, status=403)

    user_id = request.POST.get('user_id')
    new_password = request.POST.get('new_password')

    if not new_password or len(new_password) < 8:
        return JsonResponse({
            'status': 'error',
            'message': _("Password must be at least 8 characters long.")
        }, status=400)

    try:
        user = User.objects.get(id=user_id)
        user.set_password(new_password)
        user.save()

        # Log the password change
        UserActionLog.objects.create(
            user=request.user,
            action_type='change_password',
            details=f"Changed password for user '{user.username}'",
            performed_by=request.user
        )

        return JsonResponse({
            'status': 'success',
            'message': _("Password changed successfully.")
        })
    except User.DoesNotExist:
        return JsonResponse({'status': 'error', 'message': _("User not found.")}, status=404)
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)

@login_required
def templates_list(request):
    """Отображает список шаблонов документов"""
    user = request.user
    organization = user.organization
    
    # Получаем шаблоны, доступные пользователю
    # Это собственные шаблоны организации и публичные шаблоны других организаций
    templates = DocumentTemplate.objects.filter(
        Q(organization=organization) | 
        Q(is_public=True)
    ).order_by('-created_at')
    
    # Фильтры
    category_filter = request.GET.get('category', '')
    search_query = request.GET.get('search', '')
    
    if category_filter:
        templates = templates.filter(category=category_filter)
    
    if search_query:
        templates = templates.filter(
            Q(name__icontains=search_query) | 
            Q(description__icontains=search_query)
        )
    
    # Пагинация
    paginator = Paginator(templates, 12)
    page_number = request.GET.get('page')
    try:
        page_obj = paginator.page(page_number)
    except PageNotAnInteger:
        page_obj = paginator.page(1)
    except EmptyPage:
        page_obj = paginator.page(paginator.num_pages)
    
    context = {
        'templates': page_obj,
        'categories': DocumentTemplate.TEMPLATE_CATEGORIES,
        'current_category': category_filter,
        'search_query': search_query,
        'can_create': user.role in ['admin', 'manager'],
    }
    
    return render(request, 'staffs/templates_list.html', context)


@login_required
def template_detail(request, template_id):
    """Отображает детали шаблона и позволяет создать документ на его основе"""
    template = get_object_or_404(DocumentTemplate, id=template_id)
    user = request.user
    organization = user.organization
    
    # Проверка доступа
    if template.organization != organization and not template.is_public:
        django_messages.error(request, _("У вас нет доступа к этому шаблону."))
        return redirect('staffs:templates_list')
    
    if request.method == 'POST':
        # Создание документа на основе шаблона
        document = Document(
            document_name=request.POST.get('document_name'),
            document_description=request.POST.get('document_description'),
            sender=user.userprofile,
            sender_organization=organization,
            template=template,
            category=template.category,
            status='draft'
        )
        
        # Копируем файл шаблона
        if template.template_file:
            # Получаем расширение файла
            extension = os.path.splitext(template.template_file.path)[1]
            # Создаем временный файл
            with tempfile.NamedTemporaryFile(suffix=extension, delete=False) as temp_file:
                temp_file.write(template.template_file.read())
                temp_path = temp_file.name
            
            # Сохраняем в поле документа
            with open(temp_path, 'rb') as f:
                filename = f"{document.document_name}{extension}"
                document.document_content.save(filename, File(f))
            
            # Удаляем временный файл
            os.unlink(temp_path)
        
        document.save()
        
        # Создаем первую версию документа
        if document.document_content:
            version = DocumentVersion(
                document=document,
                version_number=1,
                created_by=user,
                comment="Первая версия на основе шаблона"
            )
            # Копируем файл документа в версию
            with open(document.document_content.path, 'rb') as f:
                version.content.save(os.path.basename(document.document_content.name), File(f))
            version.save()
        
        # Логируем создание документа
        UserActionLog.objects.create(
            user=user,
            action_type='create_from_template',
            details=f"Создан документ '{document.document_name}' на основе шаблона '{template.name}'",
            performed_by=user
        )
        
        django_messages.success(request, _("Документ успешно создан на основе шаблона."))
        return redirect('staffs:document-detail', document_id=document.id)
    
    context = {
        'template': template,
        'can_edit': template.organization == organization and user.role in ['admin', 'manager'],
    }
    
    return render(request, 'staffs/template_detail.html', context)


@login_required
def template_create(request):
    """Создание нового шаблона документа"""
    user = request.user
    
    # Проверка прав
    if user.role not in ['admin', 'manager']:
        django_messages.error(request, _("У вас нет прав на создание шаблонов."))
        return redirect('staffs:templates_list')
    
    if request.method == 'POST':
        name = request.POST.get('name')
        category = request.POST.get('category')
        description = request.POST.get('description')
        is_public = 'is_public' in request.POST
        template_file = request.FILES.get('template_file')
        
        if not name or not template_file:
            django_messages.error(request, _("Название и файл шаблона обязательны для заполнения."))
        else:
            template = DocumentTemplate(
                name=name,
                category=category,
                description=description,
                is_public=is_public,
                created_by=user,
                organization=user.organization,
                template_file=template_file
            )
            template.save()
            
            # Логируем создание шаблона
            UserActionLog.objects.create(
                user=user,
                action_type='create_template',
                details=f"Создан шаблон документа '{template.name}'",
                performed_by=user
            )
            
            django_messages.success(request, _("Шаблон документа успешно создан."))
            return redirect('staffs:template_detail', template_id=template.id)
    
    context = {
        'categories': DocumentTemplate.TEMPLATE_CATEGORIES,
    }
    
    return render(request, 'staffs/template_create.html', context)


@login_required
def template_edit(request, template_id):
    """Редактирование шаблона документа"""
    user = request.user
    template = get_object_or_404(DocumentTemplate, id=template_id)
    
    # Проверка прав
    if template.organization != user.organization or user.role not in ['admin', 'manager']:
        django_messages.error(request, _("У вас нет прав на редактирование этого шаблона."))
        return redirect('staffs:templates_list')
    
    if request.method == 'POST':
        name = request.POST.get('name')
        category = request.POST.get('category')
        description = request.POST.get('description')
        is_public = 'is_public' in request.POST
        template_file = request.FILES.get('template_file')
        
        if not name:
            django_messages.error(request, _("Название шаблона обязательно для заполнения."))
        else:
            template.name = name
            template.category = category
            template.description = description
            template.is_public = is_public
            
            if template_file:
                template.template_file = template_file
                
            template.save()
            
            # Логируем редактирование шаблона
            UserActionLog.objects.create(
                user=user,
                action_type='edit_template',
                details=f"Отредактирован шаблон документа '{template.name}'",
                performed_by=user
            )
            
            django_messages.success(request, _("Шаблон документа успешно обновлен."))
            return redirect('staffs:template_detail', template_id=template.id)
    
    context = {
        'template': template,
        'categories': DocumentTemplate.TEMPLATE_CATEGORIES,
    }
    
    return render(request, 'staffs/template_edit.html', context)


@login_required
@require_POST
def template_delete(request, template_id):
    """Удаление шаблона документа"""
    user = request.user
    template = get_object_or_404(DocumentTemplate, id=template_id)
    
    # Проверка прав
    if template.organization != user.organization or user.role not in ['admin', 'manager']:
        django_messages.error(request, _("У вас нет прав на удаление этого шаблона."))
        return redirect('staffs:templates_list')
    
    template_name = template.name
    template.delete()
    
    # Логируем удаление шаблона
    UserActionLog.objects.create(
        user=user,
        action_type='delete_template',
        details=f"Удален шаблон документа '{template_name}'",
        performed_by=user
    )
    
    django_messages.success(request, _("Шаблон документа успешно удален."))
    return redirect('staffs:templates_list')

@login_required
def document_versions(request, document_id):
    """Отображает список версий документа"""
    document = get_object_or_404(Document, id=document_id)
    user = request.user
    organization = user.organization
    
    # Проверка доступа к документу
    has_access = False
    if organization.is_prime_tech:
        has_access = (
            document.sender_organization == organization or
            document.recipient_organization == organization
        )
    else:
        has_access = (
            (document.sender == user.userprofile or document.recipient == user) and
            (document.sender_organization == organization or document.recipient_organization == organization)
        )
    
    if not has_access:
        django_messages.error(request, _("У вас нет доступа к этому документу."))
        return redirect('staffs:dashboard')
    
    # Получаем все версии документа
    versions = document.versions.all().order_by('-version_number')
    
    context = {
        'document': document,
        'versions': versions,
        'can_add_version': document.sender == user.userprofile or (organization.is_prime_tech and user.role in ['admin', 'manager']),
    }
    
    return render(request, 'staffs/document_versions.html', context)


@login_required
def version_detail(request, document_id, version_number):
    """Отображает детали версии документа"""
    document = get_object_or_404(Document, id=document_id)
    user = request.user
    organization = user.organization
    
    # Проверка доступа к документу
    has_access = False
    if organization.is_prime_tech:
        has_access = (
            document.sender_organization == organization or
            document.recipient_organization == organization
        )
    else:
        has_access = (
            (document.sender == user.userprofile or document.recipient == user) and
            (document.sender_organization == organization or document.recipient_organization == organization)
        )
    
    if not has_access:
        django_messages.error(request, ("У вас нет доступа к этому документу."))
        return redirect('staffs:dashboard')
    
    # Получаем версию
    try:
        version = DocumentVersion.objects.get(document=document, version_number=version_number)
    except DocumentVersion.DoesNotExist:
        django_messages.error(request, ("Запрашиваемая версия документа не найдена."))
        return redirect('staffs:document_versions', document_id=document_id)
    
    # Получаем подписи этой версии
    signatures = version.signatures.all().select_related('signer', 'certificate')
    
    # Формируем контекст для отображения предпросмотра
    page_images_base64 = []
    if version.content:
        file_path = version.content.path
        if os.path.exists(file_path):
            content_type, _ = mimetypes.guess_type(file_path)
            try:
                if content_type == 'application/pdf':
                    images = convert_from_path(file_path, dpi=200)
                    for img in images:
                        buffered = io.BytesIO()
                        img.save(buffered, format="PNG")
                        img_base64 = base64.b64encode(buffered.getvalue()).decode('utf-8')
                        page_images_base64.append(img_base64)
                elif content_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 
                                     'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                                     'text/plain']:
                    # Используем существующую функцию text_to_images для предпросмотра
                    try:
                        if content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                            doc = DocxDocument(file_path)
                            text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
                        elif content_type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet':
                            df = pd.read_excel(file_path)
                            text = df.to_string(index=False)
                        else:  # text/plain
                            with open(file_path, 'r', encoding='utf-8') as f:
                                text = f.read()
                        
                        images = text_to_images(text, width=800, height=1200)
                        for img in images:
                            buffered = io.BytesIO()
                            img.save(buffered, format="PNG")
                            img_base64 = base64.b64encode(buffered.getvalue()).decode('utf-8')
                            page_images_base64.append(img_base64)
                    except Exception as e:
                        logger.error(f"Error generating preview: {str(e)}")
                        django_messages.warning(request, _("Не удалось создать предпросмотр для этого типа файла."))
                else:
                    django_messages.warning(request, _("Предпросмотр недоступен для этого типа файла."))
            except Exception as e:
                logger.error(f"Error generating preview: {str(e)}")
                django_messages.error(request, _("Ошибка при создании предпросмотра: ") + str(e))
    
    context = {
        'document': document,
        'version': version,
        'signatures': signatures,
        'can_sign': user.certificates.filter(certificate_type='user').exists(),
        'can_restore': document.sender == user.userprofile,
        'page_images': page_images_base64,
    }
    
    return render(request, 'staffs/version_detail.html', context)


@login_required
def version_preview(request, document_id, version_number):
    """
    API endpoint для получения превью версии документа для сравнения
    Возвращает JSON с данными для отображения
    """
    document = get_object_or_404(Document, id=document_id)
    user = request.user
    organization = user.organization
    
    # Проверка доступа к документу
    has_access = False
    if organization.is_prime_tech:
        has_access = (
            document.sender_organization == organization or
            document.recipient_organization == organization
        )
    else:
        has_access = (
            (document.sender == user.userprofile or document.recipient == user) and
            (document.sender_organization == organization or document.recipient_organization == organization)
        )
    
    if not has_access:
        return JsonResponse({
            'status': 'error',
            'message': ("У вас нет доступа к этому документу.")
        })
    
    # Получаем версию
    try:
        version = DocumentVersion.objects.get(document=document, version_number=version_number)
    except DocumentVersion.DoesNotExist:
        return JsonResponse({
            'status': 'error',
            'message': ("Запрашиваемая версия документа не найдена.")
        })
    
    # Формируем данные для JSON-ответа
    page_images_base64 = []
    text_content = None
    html_content = None
    
    if version.content:
        file_path = version.content.path
        if os.path.exists(file_path):
            content_type, _ = mimetypes.guess_type(file_path)
            try:
                # Обработка PDF файлов
                if content_type == 'application/pdf':
                    images = convert_from_path(file_path, dpi=200)
                    for img in images:
                        buffered = io.BytesIO()
                        img.save(buffered, format="PNG")
                        img_base64 = base64.b64encode(buffered.getvalue()).decode('utf-8')
                        page_images_base64.append(img_base64)
                
                # Обработка DOCX/DOC файлов - конвертируем в PDF с сохранением форматирования
                elif content_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword']:
                    # Создаем временный файл для PDF
                    with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_pdf:
                        temp_pdf_path = temp_pdf.name
                    
                    try:
                        # Используем улучшенную функцию конвертации
                        conversion_success = convert_docx_to_pdf(file_path, temp_pdf_path)
                        
                        # Если конвертация удалась, генерируем изображения страниц
                        if conversion_success:
                            # Используем более высокое разрешение для лучшего качества
                            images = convert_from_path(temp_pdf_path, dpi=300)
                            for img in images:
                                buffered = io.BytesIO()
                                img.save(buffered, format="PNG", quality=95)  # Высокое качество
                                img_base64 = base64.b64encode(buffered.getvalue()).decode('utf-8')
                                page_images_base64.append(img_base64)
                        else:
                            # Если конвертация не удалась, используем резервный метод text_to_images
                            logger.warning(f"Failed to convert DOCX to PDF, using fallback method")
                            doc = DocxDocument(file_path)
                            text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
                            if not text:
                                logger.warning(f"DOCX file is empty: {file_path}")
                                django_messages.warning(request, ("The DOCX file is empty or contains no readable text."))
                            else:
                                images = text_to_images(text, width=800, height=1200)
                                for img in images:
                                    buffered = io.BytesIO()
                                    img.save(buffered, format="PNG")
                                    img_base64 = base64.b64encode(buffered.getvalue()).decode('utf-8')
                                    page_images_base64.append(img_base64)
                    finally:
                        # Удаляем временный PDF файл
                        if os.path.exists(temp_pdf_path):
                            os.unlink(temp_pdf_path)
                
                elif content_type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet':
                    df = pd.read_excel(file_path)
                    text_content = df.to_string(index=False)
                    
                    # Создаем HTML таблицу с улучшенным форматированием
                    html_content = '<div class="overflow-x-auto">'
                    html_content += df.to_html(
                        classes='min-w-full divide-y divide-gray-200 border', 
                        index=False,
                        justify='center',
                        bold_rows=True,
                        border=1
                    ).replace('<table', '<table class="min-w-full divide-y divide-gray-200"')
                    html_content += '</div>'
                
                elif content_type == 'text/plain':
                    with open(file_path, 'r', encoding='utf-8') as f:
                        text_content = f.read()
                    
                    # Создаем HTML структуру для текста
                    paragraphs = text_content.split('\n')
                    html_paragraphs = []
                    for p in paragraphs:
                        if p.strip():
                            html_paragraphs.append(f'<p class="my-2">{p}</p>')
                        else:
                            html_paragraphs.append('<div class="h-4"></div>')  # Пустая строка как отступ
                    
                    html_content = f'<div class="document-content prose max-w-none">{" ".join(html_paragraphs)}</div>'
                
                else:
                    return JsonResponse({
                        'status': 'warning',
                        'message': _("Предпросмотр недоступен для этого типа файла.")
                    })
            
            except Exception as e:
                logger.error(f"Ошибка при создании предпросмотра для {file_path}: {str(e)}", exc_info=True)
                return JsonResponse({
                    'status': 'error',
                    'message': f"{_('Ошибка при создании предпросмотра:')} {str(e)}"
                })
        else:
            return JsonResponse({
                'status': 'error',
                'message': ("Файл не найден на сервере.")
            })
    else:
        return JsonResponse({
            'status': 'warning',
            'message': ("Нет файла, прикрепленного к этой версии.")
        })

    return JsonResponse({
        'status': 'success',
        'images': page_images_base64,
        'text_content': text_content,
        'html_content': html_content,
        'file_name': os.path.basename(version.content.name) if version.content else "",
        'version': version.version_number
    })


@login_required
def sign_document(request, document_id, version_number=None):
    """Подписывает документ или конкретную версию документа"""
    document = get_object_or_404(Document, id=document_id)
    user = request.user
    organization = user.organization
    
    # Проверка доступа
    has_access = False
    if organization.is_prime_tech:
        has_access = (
            document.sender_organization == organization or
            document.recipient_organization == organization
        )
    else:
        has_access = (
            (document.sender == user.userprofile or document.recipient == user) and
            (document.sender_organization == organization or document.recipient_organization == organization)
        )
    
    if not has_access:
        django_messages.error(request, _("У вас нет доступа к этому документу."))
        return redirect('staffs:dashboard')
    
    if request.method == 'POST':
        form = SendDocumentForm(request.POST, request.FILES, user=request.user)
        if form.is_valid():
            document = form.save(commit=False)
            document.sender = request.user.userprofile
            document.sender_organization = request.user.organization
            
            # Автоматически устанавливаем организацию получателя
            recipient_user = document.recipient
            if recipient_user:
                document.recipient_organization = recipient_user.organization
            
            # Извлекаем информацию из документа
            if document.document_content:
                try:
                    # Определяем тип файла
                    file_path = document.document_content.path
                    content_type = mimetypes.guess_type(file_path)[0]

                    # Извлекаем текст и информацию в зависимости от типа файла
                    if content_type == 'application/pdf':
                        reader = PdfReader(file_path)
                        # Получаем количество страниц
                        document.page_count = len(reader.pages)
                        # Получаем текст из первой страницы для краткого описания
                        first_page = reader.pages[0]
                        text = first_page.extract_text()
                        # Берем первые 200 символов для краткого описания
                        document.summary = text[:200] if text else None

                    elif content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                        doc = DocxDocument(file_path)
                        # Получаем количество параграфов как страницы
                        document.page_count = len(doc.paragraphs)
                        # Получаем текст из первого параграфа для краткого описания
                        text = doc.paragraphs[0].text if doc.paragraphs else None
                        document.summary = text[:200] if text else None

                    elif content_type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet':
                        df = pd.read_excel(file_path)
                        # Считаем количество листов как страницы
                        wb = openpyxl.load_workbook(file_path)
                        document.page_count = len(wb.sheetnames)
                        # Получаем первые несколько строк для краткого описания
                        text = df.head().to_string()
                        document.summary = text[:200] if text else None

                    elif content_type == 'text/plain':
                        with open(file_path, 'r', encoding='utf-8') as f:
                            text = f.read()
                            # Считаем количество строк как страницы
                            document.page_count = len(text.splitlines())
                            # Берем первые 200 символов для краткого описания
                            document.summary = text[:200] if text else None

                    # Устанавливаем имя документа, если оно не задано
                    if not document.document_name:
                        document.document_name = os.path.basename(file_path)

                    # Устанавливаем метод отправки по умолчанию
                    if not document.method:
                        document.method = 'Внутренняя система'

                except Exception as e:
                    logger.error(f"Error extracting document info: {str(e)}")
                    # В случае ошибки устанавливаем базовые значения
                    if not document.summary:
                        document.summary = _("Document content could not be extracted")
                    if not document.page_count:
                        document.page_count = 1

            document.status = 'sent'
            document.date_sent = timezone.now()
            document.save()

            # Логирование действия отправки документа
            UserActionLog.objects.create(
                user=request.user,
                action_type='send_document',
                details=f"Sent document '{document.document_name}' to {document.recipient.username} ({document.recipient_organization.name if document.recipient_organization else 'No organization'})",
                performed_by=request.user
            )

            # Создаём уведомление для получателя
            if document.recipient:
                Notification.objects.create(
                    user=document.recipient,
                    message=f"New document '{document.document_name}' received from {document.sender.user.username} ({document.sender_organization.name})"
                )

            django_messages.success(request, _("Document sent successfully."))
            return redirect('staffs:dashboard')
        else:
            for field, errors in form.errors.items():
                for error in errors:
                    django_messages.error(request, f"{field}: {error}")
    else:
        form = SendDocumentForm(user=request.user)
    return render(request, 'staffs/send.html', {'form': form})


@login_required
def verify_signature(request, signature_id):
    """Проверяет действительность подписи"""
    signature = get_object_or_404(DigitalSignature, id=signature_id)
    
    # Здесь будет реальная логика проверки подписи
    # В данном примере просто обновляем дату проверки
    
    signature.validation_date = timezone.now()
    signature.save()
    
    # Логируем проверку подписи
    UserActionLog.objects.create(
        user=request.user,
        action_type='verify_signature',
        details=f"Проверена подпись документа '{signature.document.document_name}' от пользователя {signature.signer.username}",
        performed_by=request.user
    )
    
    status = "действительна" if signature.is_valid else "недействительна"
    django_messages.info(request, _(f"Подпись {status}. Последняя проверка: {signature.validation_date}"))
    
    # Возвращаемся на страницу, с которой был запрос
    referer = request.META.get('HTTP_REFERER')
    if referer:
        return redirect(referer)
    else:
        return redirect('staffs:document_detail', document_id=signature.document.id)

@login_required
def create_version(request, document_id):
    """Создание новой версии документа"""
    document = get_object_or_404(Document, id=document_id)
    user = request.user
    organization = user.organization
    
    # Проверка доступа и прав на создание новой версии
    has_access = False
    if document.sender == user.userprofile or (organization.is_prime_tech and user.role in ['admin', 'manager']):
        has_access = True
    
    if not has_access:
        django_messages.error(request, _("У вас нет прав на создание новой версии этого документа."))
        return redirect('staffs:document-detail', document_id=document_id)
    
    if request.method == 'POST':
        try:
            # Получаем новый файл и комментарий
            new_file = request.FILES.get('document_file')
            comment = request.POST.get('comment', '')
            
            if not new_file:
                django_messages.error(request, _("Необходимо предоставить файл для новой версии."))
                return render(request, 'staffs/create_version.html', {'document': document})
            
            # Получаем последнюю версию документа
            latest_version = document.versions.order_by('-version_number').first()
            new_version_number = 1
            if latest_version:
                new_version_number = latest_version.version_number + 1
            
            # Создаем временную копию файла
            with tempfile.NamedTemporaryFile(delete=False) as temp_file:
                for chunk in new_file.chunks():
                    temp_file.write(chunk)
                temp_path = temp_file.name
            
            # Создаем новую версию, используя сохраненный временный файл
            with open(temp_path, 'rb') as f:
                version = DocumentVersion(
                    document=document,
                    version_number=new_version_number,
                    created_by=user,
                    comment=comment
                )
                # Явно сохраняем содержимое в поле файла
                filename = os.path.basename(new_file.name)
                version.content.save(filename, File(f))
                version.save()
                
                # Перематываем файл в начало для повторного использования
                f.seek(0)
                # Обновляем основной файл документа текущей версией
                document.document_content.save(filename, File(f))
            
            # Удаляем временный файл
            os.unlink(temp_path)
            
            # Сохраняем документ после обновления файла
            document.save()
            
            # Логируем создание новой версии
            UserActionLog.objects.create(
                user=user,
                action_type='create_version',
                details=f"Создана новая версия {new_version_number} для документа '{document.document_name}'",
                performed_by=user
            )
            
            django_messages.success(request, _(f"Версия {new_version_number} успешно создана."))
            return redirect('staffs:document_versions', document_id=document_id)
            
        except Exception as e:
            logger.error(f"Ошибка при создании версии документа: {str(e)}", exc_info=True)
            django_messages.error(request, _(f"Произошла ошибка при создании версии: {str(e)}"))
            return render(request, 'staffs/create_version.html', {'document': document})
    
    return render(request, 'staffs/create_version.html', {'document': document})


@login_required
def restore_version(request, document_id, version_number):
    """Восстановление предыдущей версии документа"""
    document = get_object_or_404(Document, id=document_id)
    user = request.user
    
    # Проверка доступа
    if document.sender != user.userprofile and not (user.organization.is_prime_tech and user.role in ['admin', 'manager']):
        django_messages.error(request, _("У вас нет прав на восстановление версий этого документа."))
        return redirect('staffs:document_versions', document_id=document_id)
    
    try:
        # Получаем версию, которую нужно восстановить
        old_version = DocumentVersion.objects.get(document=document, version_number=version_number)
        
        # Получаем последнюю версию
        latest_version = document.versions.order_by('-version_number').first()
        new_version_number = latest_version.version_number + 1 if latest_version else 1
        
        # Создаем временную копию файла
        with tempfile.NamedTemporaryFile(delete=False) as temp_file:
            with old_version.content.open('rb') as content_file:
                temp_file.write(content_file.read())
            temp_path = temp_file.name
        
        # Создаем новую версию на основе старой
        with open(temp_path, 'rb') as f:
            # Создаем новую версию
            new_version = DocumentVersion(
                document=document,
                version_number=new_version_number,
                created_by=user,
                comment=f"Восстановлено из версии {version_number}"
            )
            
            filename = os.path.basename(old_version.content.name)
            new_version.content.save(filename, File(f))
            new_version.save()
            
            # Перематываем файл в начало для повторного использования
            f.seek(0)
            # Обновляем основной файл документа
            document.document_content.save(filename, File(f))
        
        # Удаляем временный файл
        os.unlink(temp_path)
        
        # Логируем восстановление версии
        UserActionLog.objects.create(
            user=user,
            action_type='restore_version',
            details=f"Восстановлена версия {version_number} для документа '{document.document_name}'",
            performed_by=user
        )
        
        django_messages.success(request, _(f"Версия {version_number} успешно восстановлена как версия {new_version_number}."))
    except DocumentVersion.DoesNotExist:
        django_messages.error(request, _("Указанная версия не найдена."))
    except Exception as e:
        logger.error(f"Ошибка при восстановлении версии документа: {str(e)}", exc_info=True)
        django_messages.error(request, _(f"Произошла ошибка при восстановлении версии: {str(e)}"))
    
    return redirect('staffs:document_versions', document_id=document_id)

# Добавляем функцию определения типа SQL файла
def determine_sql_file_type(file_path):
    """
    Определяет тип SQL файла: SQLite, PostgreSQL или Generic SQL
    
    Args:
        file_path: путь к SQL файлу
        
    Returns:
        Строка с типом: "SQLite", "PostgreSQL", "PostgreSQL (15)", "SQL (Generic)" или "Unknown"
    """
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read(4000)  # Читаем больше контента для более точного определения
            
            # SQLite маркеры
            sqlite_markers = [
                'BEGIN TRANSACTION;',
                'PRAGMA foreign_keys=',
                'sqlite_sequence',
                '.schema',
                'CREATE TABLE sqlite_',
                'COMMIT;'  # Очень часто используется в SQLite дампах
            ]
            
            # PostgreSQL маркеры
            postgres_markers = [
                'pg_dump',
                'PostgreSQL database dump',
                'SET statement_timeout',
                'SET lock_timeout',
                'SET client_encoding',
                'SET standard_conforming_strings',
                'CREATE EXTENSION',
                'ALTER TABLE ONLY',
                'SCHEMA public',
                'pg_catalog'
            ]
            
            # Проверка версии PostgreSQL
            pg_version = None
            pg_version_matches = [
                r'PostgreSQL database dump.*(\d+)\.(\d+)',  # Ищем номер версии
                r'pg_dump.*(\d+)\.(\d+)',  # Ищем номер версии в выводе pg_dump
            ]
            
            for pattern in pg_version_matches:
                import re
                match = re.search(pattern, content)
                if match:
                    major_version = match.group(1)
                    pg_version = major_version
                    break
            
            # Проверяем наличие маркеров SQLite - любой из них подтверждает тип
            for marker in sqlite_markers:
                if marker in content:
                    return "SQLite"
            
            # Проверяем наличие маркеров PostgreSQL
            for marker in postgres_markers:
                if marker in content:
                    if pg_version:
                        return f"PostgreSQL ({pg_version})"
                    return "PostgreSQL"
            
            # Проверяем общие SQL команды
            if ('CREATE TABLE' in content or 
                'INSERT INTO' in content or 
                'ALTER TABLE' in content or
                'SELECT ' in content):
                # Если файл содержит типичные структуры SQLite
                if 'integer primary key autoincrement' in content.lower():
                    return "SQLite"
                # Если файл содержит типичные структуры PostgreSQL
                elif 'serial primary key' in content.lower() or 'bigserial' in content.lower():
                    return "PostgreSQL"
                return "SQL (Generic)"
                
            return "Unknown"
    except Exception as e:
        print(f"Error determining SQL file type: {str(e)}")
        return "Unknown"
